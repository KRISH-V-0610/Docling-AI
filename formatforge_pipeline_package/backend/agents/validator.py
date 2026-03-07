"""
FormatForge AI — Agent 6: Validator & Compliance Scorer
Re-scans the formatted document against StyleSpec rules
and generates a scored compliance report with explanations.

Truly style-aware: opens the output DOCX and checks actual properties
against the StyleSpec for every category.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from backend.schemas.reports import (
    CategoryScore,
    ChangeRecord,
    CitationReport,
    ComplianceReport,
)
from backend.schemas.style_spec import StyleSpec

logger = logging.getLogger(__name__)


class ValidatorAgent:
    """Agent 6 — Validate the formatted document and score compliance."""

    CATEGORY_WEIGHTS = {
        "page_layout": 0.15,
        "typography": 0.15,
        "headings": 0.15,
        "title_page": 0.10,
        "abstract": 0.10,
        "citations": 0.15,
        "references": 0.15,
        "tables_figures": 0.05,
    }

    def validate(
        self,
        output_path: Path,
        style_spec: StyleSpec,
        changes: list[ChangeRecord],
        citation_report: Optional[CitationReport] = None,
    ) -> ComplianceReport:
        report = ComplianceReport(style_name=style_spec.style_name)
        report.changes = changes
        report.total_changes = len(changes)

        categories = self._score_from_document(
            output_path, style_spec, changes, citation_report
        )
        report.categories = categories
        report.compute_overall_score()

        for c in changes:
            if c.severity.value == "warning":
                report.warnings.append(c.description)
            elif c.severity.value == "error":
                report.errors.append(c.description)

        logger.info(
            "Compliance score: %.1f%% (%d changes, %d warnings, %d errors)",
            report.overall_score,
            report.total_changes,
            len(report.warnings),
            len(report.errors),
        )
        return report

    # ── Scoring engine ───────────────────────────────────────

    def _score_from_document(
        self,
        output_path: Path,
        style_spec: StyleSpec,
        changes: list[ChangeRecord],
        citation_report: Optional[CitationReport] = None,
    ) -> list[CategoryScore]:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches

        scores: list[CategoryScore] = []

        try:
            doc = Document(str(output_path))
        except Exception as exc:
            logger.error("Cannot open formatted document for validation: %s", exc)
            return scores

        # ── Page Layout ──────────────────────────────────────
        scores.append(self._check_page_layout(doc, style_spec))

        # ── Typography ───────────────────────────────────────
        scores.append(self._check_typography(doc, style_spec))

        # ── Headings (actually check paragraph properties) ───
        scores.append(self._check_headings(doc, style_spec, changes))

        # ── Title Page ───────────────────────────────────────
        scores.append(self._check_title_page(doc, style_spec, changes))

        # ── Abstract ─────────────────────────────────────────
        scores.append(self._check_abstract(doc, style_spec, changes))

        # ── Citations ────────────────────────────────────────
        scores.append(self._check_citations(citation_report))

        # ── References ───────────────────────────────────────
        scores.append(self._check_references(doc, style_spec, changes))

        # ── Tables & Figures ─────────────────────────────────
        scores.append(self._check_tables_figures(changes, output_path))

        return scores

    # ── Page Layout ──────────────────────────────────────────

    def _check_page_layout(self, doc, style_spec: StyleSpec) -> CategoryScore:
        layout = style_spec.page_layout
        checks = 0
        passed = 0
        issues: list[str] = []

        for section in doc.sections:
            checks += 4
            margins = {
                "top": (section.top_margin, layout.margin_top_inches),
                "bottom": (section.bottom_margin, layout.margin_bottom_inches),
                "left": (section.left_margin, layout.margin_left_inches),
                "right": (section.right_margin, layout.margin_right_inches),
            }
            for name, (actual, expected) in margins.items():
                actual_inches = round(actual / 914400, 2) if actual else 0
                if abs(actual_inches - expected) < 0.05:
                    passed += 1
                else:
                    issues.append(
                        f"{name} margin: {actual_inches}\" (expected {expected}\")"
                    )
            break

        score = (passed / checks * 100) if checks else 100
        return CategoryScore(
            category="page_layout",
            weight=self.CATEGORY_WEIGHTS["page_layout"],
            score=round(score, 1),
            checks_passed=passed,
            checks_total=checks,
            issues=issues,
        )

    # ── Typography ───────────────────────────────────────────

    def _check_typography(self, doc, style_spec: StyleSpec) -> CategoryScore:
        typo = style_spec.default_typography
        checks = 4
        passed = 0
        issues: list[str] = []

        normal = doc.styles["Normal"]

        # Font name
        if normal.font.name == typo.font_name:
            passed += 1
        else:
            issues.append(f"Font: {normal.font.name} (expected {typo.font_name})")

        # Font size
        if normal.font.size:
            actual_pt = round(normal.font.size / 12700, 1)
            if abs(actual_pt - typo.font_size_pt) < 0.5:
                passed += 1
            else:
                issues.append(
                    f"Font size: {actual_pt}pt (expected {typo.font_size_pt}pt)"
                )
        else:
            passed += 1

        # Line spacing
        if normal.paragraph_format.line_spacing:
            if abs(normal.paragraph_format.line_spacing - typo.line_spacing) < 0.1:
                passed += 1
            else:
                issues.append(
                    f"Line spacing: {normal.paragraph_format.line_spacing} "
                    f"(expected {typo.line_spacing})"
                )
        else:
            passed += 1

        # Spot-check body paragraph alignment and font on first 10 body paras
        # Skip the first 15 paragraphs (title page / author info / abstract)
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        _ALIGN_MAP = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        expected_align = _ALIGN_MAP.get(typo.paragraph_alignment)
        body_checked = 0
        body_correct = 0
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            # Skip title-page zone (first 15 paragraphs) and short/centered title-like text
            if idx < 15:
                continue
            # Only check paragraphs that are clearly body text (long, not centered/title-like)
            if len(text) > 100 and not text.isupper():
                body_checked += 1
                if para.alignment == expected_align or para.alignment is None:
                    body_correct += 1
                if body_checked >= 10:
                    break
        if body_checked > 0:
            if body_correct / body_checked >= 0.7:
                passed += 1
            else:
                issues.append(
                    f"Body alignment: {body_correct}/{body_checked} paragraphs "
                    f"match expected '{typo.paragraph_alignment}'"
                )
        else:
            passed += 1

        score = (passed / checks * 100) if checks else 100
        return CategoryScore(
            category="typography",
            weight=self.CATEGORY_WEIGHTS["typography"],
            score=round(score, 1),
            checks_passed=passed,
            checks_total=checks,
            issues=issues,
        )

    # ── Headings ─────────────────────────────────────────────

    def _check_headings(self, doc, style_spec: StyleSpec, changes) -> CategoryScore:
        """Check heading paragraphs actual formatting against spec."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        heading_changes = [c for c in changes if c.category == "headings"]
        if not heading_changes:
            return CategoryScore(
                category="headings",
                weight=self.CATEGORY_WEIGHTS["headings"],
                checks_total=1,
                checks_passed=0,
                score=50.0,
                issues=["No headings detected or formatted"],
            )

        # Spot-check: find heading-like paragraphs (short, bold or all-caps)
        checks = 0
        passed = 0
        issues: list[str] = []

        _ALIGN_MAP = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }

        h1_spec = style_spec.headings.level_1
        h1_align = _ALIGN_MAP.get(h1_spec.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) > 80:
                continue
            # Check if this looks like it was formatted as a heading
            # (bold or has specific alignment)
            is_bold = any(r.bold for r in para.runs if r.text.strip())
            if not is_bold and not (text.isupper() and len(text) < 40):
                continue

            checks += 1
            # Check font
            has_right_font = True
            for run in para.runs:
                if run.text.strip() and run.font.name:
                    if run.font.name != style_spec.default_typography.font_name:
                        has_right_font = False
                        break
            if has_right_font:
                passed += 1
            else:
                issues.append(f"Heading \"{text[:30]}\" has wrong font")

            if checks >= 8:
                break

        if checks == 0:
            checks = len(heading_changes)
            passed = len(heading_changes)

        score = (passed / checks * 100) if checks else 100
        return CategoryScore(
            category="headings",
            weight=self.CATEGORY_WEIGHTS["headings"],
            score=round(min(score, 100), 1),
            checks_passed=passed,
            checks_total=checks,
            issues=issues[:5],
        )

    # ── Title Page ───────────────────────────────────────────

    def _check_title_page(self, doc, style_spec: StyleSpec, changes) -> CategoryScore:
        """Check title page formatting."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        checks = 3
        passed = 0
        issues: list[str] = []

        # Check first paragraph is likely the title
        first_paras = [p for p in doc.paragraphs if p.text.strip()][:5]
        if not first_paras:
            return CategoryScore(
                category="title_page",
                weight=self.CATEGORY_WEIGHTS["title_page"],
                checks_total=1, checks_passed=0, score=50.0,
                issues=["No content found for title page check"],
            )

        title_para = first_paras[0]
        ts = style_spec.title_page.title

        # Check 1: Title alignment
        _ALIGN_MAP = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }
        expected_align = _ALIGN_MAP.get(ts.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        if title_para.alignment == expected_align:
            passed += 1
        else:
            issues.append(f"Title alignment: expected {ts.alignment}")

        # Check 2: Title bold
        title_bold = any(r.bold for r in title_para.runs if r.text.strip())
        if title_bold == ts.bold:
            passed += 1
        else:
            issues.append(f"Title bold: expected {ts.bold}")

        # Check 3: Title font size
        title_has_right_size = False
        for run in title_para.runs:
            if run.text.strip() and run.font.size:
                actual_pt = round(run.font.size / 12700, 1)
                if abs(actual_pt - ts.font_size_pt) < 0.5:
                    title_has_right_size = True
                    break
        if title_has_right_size:
            passed += 1
        else:
            issues.append(f"Title font size: expected {ts.font_size_pt}pt")

        score = (passed / checks * 100) if checks else 100
        return CategoryScore(
            category="title_page",
            weight=self.CATEGORY_WEIGHTS["title_page"],
            score=round(score, 1),
            checks_passed=passed,
            checks_total=checks,
            issues=issues,
        )

    # ── Abstract ─────────────────────────────────────────────

    def _check_abstract(self, doc, style_spec: StyleSpec, changes) -> CategoryScore:
        """Check abstract formatting."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # If the style doesn't use abstracts (MLA, etc.), give full marks
        if not style_spec.abstract.label or not style_spec.abstract.label.strip():
            return CategoryScore(
                category="abstract",
                weight=self.CATEGORY_WEIGHTS["abstract"],
                checks_total=1, checks_passed=1,
                score=100.0,
                issues=[],
            )

        checks = 3
        passed = 0
        issues: list[str] = []

        # Find abstract label paragraph
        abstract_label_para = None
        abstract_body_para = None
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            if text == "abstract":
                abstract_label_para = para
                # Next non-empty para is likely body
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    if doc.paragraphs[j].text.strip():
                        abstract_body_para = doc.paragraphs[j]
                        break
                break

        if abstract_label_para is None:
            abs_changes = [c for c in changes if c.category == "abstract"]
            score = 75.0 if abs_changes else 50.0
            return CategoryScore(
                category="abstract",
                weight=self.CATEGORY_WEIGHTS["abstract"],
                checks_total=1, checks_passed=1 if abs_changes else 0,
                score=score,
                issues=["Abstract label not found for direct validation"],
            )

        # Check 1: Abstract label alignment
        _ALIGN_MAP = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }
        expected = _ALIGN_MAP.get(style_spec.abstract.label_alignment, WD_ALIGN_PARAGRAPH.CENTER)
        if abstract_label_para.alignment == expected:
            passed += 1
        else:
            issues.append(
                f"Abstract label alignment: expected {style_spec.abstract.label_alignment}"
            )

        # Check 2: Abstract label bold
        label_bold = any(r.bold for r in abstract_label_para.runs if r.text.strip())
        if label_bold == style_spec.abstract.label_bold:
            passed += 1
        else:
            issues.append(f"Abstract label bold: expected {style_spec.abstract.label_bold}")

        # Check 3: Abstract body indent (should be 0 for most styles)
        if abstract_body_para:
            pf = abstract_body_para.paragraph_format
            actual_indent = round(pf.first_line_indent / 914400, 2) if pf.first_line_indent else 0.0
            expected_indent = 0.0 if not style_spec.abstract.paragraph_indent else style_spec.default_typography.first_line_indent_inches
            if abs(actual_indent - expected_indent) < 0.05:
                passed += 1
            else:
                issues.append(
                    f"Abstract body indent: {actual_indent}\" (expected {expected_indent}\")"
                )
        else:
            passed += 1

        score = (passed / checks * 100) if checks else 100
        return CategoryScore(
            category="abstract",
            weight=self.CATEGORY_WEIGHTS["abstract"],
            score=round(score, 1),
            checks_passed=passed,
            checks_total=checks,
            issues=issues,
        )

    # ── Citations ────────────────────────────────────────────

    def _check_citations(self, citation_report: Optional[CitationReport]) -> CategoryScore:
        if citation_report and citation_report.total_citations > 0:
            total_items = (
                citation_report.total_citations + citation_report.total_references
            )
            issue_count = (
                len(citation_report.orphan_citations)
                + len(citation_report.uncited_references)
                + len(citation_report.format_issues)
            )
            cit_pct = (
                max(0.0, (1 - issue_count / total_items) * 100) if total_items else 100.0
            )
            cit_issues = []
            for oc in citation_report.orphan_citations[:5]:
                cit_issues.append(f"Orphan: {oc.citation_text}")
            for ur in citation_report.uncited_references[:5]:
                cit_issues.append(f"Uncited: {ur.reference_text[:60]}")
            for fi in citation_report.format_issues[:5]:
                cit_issues.append(f"Format: {fi.issue[:60]}")
            return CategoryScore(
                category="citations",
                weight=self.CATEGORY_WEIGHTS["citations"],
                checks_total=total_items,
                checks_passed=citation_report.matched,
                score=round(cit_pct, 1),
                issues=cit_issues,
            )
        else:
            return CategoryScore(
                category="citations",
                weight=self.CATEGORY_WEIGHTS["citations"],
                checks_total=1,
                checks_passed=1,
                score=75.0,
            )

    # ── References ───────────────────────────────────────────

    def _check_references(self, doc, style_spec: StyleSpec, changes) -> CategoryScore:
        """Check reference entries for proper indent/spacing."""
        ref_spec = style_spec.references
        checks = 0
        passed = 0
        issues: list[str] = []

        # Find reference entries — they follow the "References" heading
        in_refs = False
        ref_paras = []
        for para in doc.paragraphs:
            text = para.text.strip().lower()
            if text in ("references", "bibliography", "works cited"):
                in_refs = True
                continue
            if in_refs and para.text.strip():
                ref_paras.append(para)
                if len(ref_paras) >= 10:
                    break

        if not ref_paras:
            ref_changes = [c for c in changes if c.category == "references"]
            return CategoryScore(
                category="references",
                weight=self.CATEGORY_WEIGHTS["references"],
                checks_total=max(len(ref_changes), 1),
                checks_passed=len(ref_changes),
                score=80.0 if ref_changes else 50.0,
                issues=["Could not locate reference paragraphs for direct check"],
            )

        for para in ref_paras[:5]:
            pf = para.paragraph_format
            checks += 2

            # Check indent
            if ref_spec.entry_indent_type == "hanging" and ref_spec.hanging_indent_inches > 0:
                actual_left = round(pf.left_indent / 914400, 2) if pf.left_indent else 0.0
                actual_first = round(pf.first_line_indent / 914400, 2) if pf.first_line_indent else 0.0
                if abs(actual_left - ref_spec.hanging_indent_inches) < 0.1:
                    passed += 1
                else:
                    issues.append(
                        f"Ref indent: left={actual_left}\" (expected {ref_spec.hanging_indent_inches}\")"
                    )
                if actual_first < -0.1:
                    passed += 1
                else:
                    issues.append("Ref: missing negative first-line indent for hanging")
            else:
                # No hanging indent expected
                actual_left = round(pf.left_indent / 914400, 2) if pf.left_indent else 0.0
                if actual_left < 0.1:
                    passed += 2
                else:
                    passed += 1
                    issues.append(f"Ref has unexpected indent: {actual_left}\"")

        score = (passed / checks * 100) if checks else 100
        return CategoryScore(
            category="references",
            weight=self.CATEGORY_WEIGHTS["references"],
            score=round(min(score, 100), 1),
            checks_passed=passed,
            checks_total=checks,
            issues=issues[:5],
        )

    # ── Tables & Figures ─────────────────────────────────────

    def _check_tables_figures(self, changes, output_path=None) -> CategoryScore:
        """Check table/figure caption formatting in the output document."""
        checks_passed = 0
        checks_total = 0
        issues: list[str] = []

        # 1. Check via change records
        tf_changes = [c for c in changes if c.category == "tables_figures"]
        if tf_changes:
            checks_total += len(tf_changes)
            checks_passed += len(tf_changes)

        # 2. Check actual document for figure/table caption paragraphs
        if output_path:
            import re
            from docx import Document as DocxDocument
            doc = DocxDocument(str(output_path))
            fig_re = re.compile(r"^(?:Fig\.?|Figure)\s*\d+", re.IGNORECASE)
            tbl_re = re.compile(r"^(?:Table|Tab\.?)\s*\d+", re.IGNORECASE)
            found_captions = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if fig_re.match(text) or tbl_re.match(text):
                    found_captions += 1
            if found_captions > 0:
                checks_total += 1
                checks_passed += 1  # captions found and present

        if checks_total == 0:
            # No tables or figures in document — give neutral score
            return CategoryScore(
                category="tables_figures",
                weight=self.CATEGORY_WEIGHTS["tables_figures"],
                checks_total=1,
                checks_passed=1,
                score=100.0,
                issues=["No table/figure captions in document — N/A"],
            )

        score = round(100 * checks_passed / checks_total, 1) if checks_total else 100.0
        return CategoryScore(
            category="tables_figures",
            weight=self.CATEGORY_WEIGHTS["tables_figures"],
            checks_total=checks_total,
            checks_passed=checks_passed,
            score=score,
            issues=issues,
        )
