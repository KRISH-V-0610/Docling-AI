"""
FormatForge AI — Agent 5: Transformation Engine
Takes labeled DocIR + StyleSpec → produces a publication-ready DOCX.

Phase 2 — Complete formatting pipeline:
 • Page layout (margins, page size, orientation)
 • Default typography (Normal style — font, size, spacing)
 • Per-element formatting for every semantic role
 • Font enforcement on every run (overrides PDF-to-DOCX artefacts)
 • Running head (shortened title ALL-CAPS left, page number right)
 • Title-page & abstract-page section breaks
 • Automatic label insertion ("Abstract", "References") when missing
 • Table / figure caption formatting (APA style)

100 % deterministic — zero LLM calls.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from backend.config import OUTPUT_FORMATTED_DIR
from backend.schemas.docir import DocElement, DocIR, ElementRole, ElementType
from backend.schemas.reports import ChangeRecord, ChangeStatus, Severity
from backend.schemas.style_spec import StyleSpec
from backend.agents.paragraph_merger import ParagraphMerger

# Module-level alignment map (avoids repeated dict creation)
_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

logger = logging.getLogger(__name__)

# ── Regex to strip existing heading numbering ────────────────
#    Matches: "I. ", "II. ", "A. ", "1. ", "1) ", "a) ", etc.
_EXISTING_NUM_RE = re.compile(
    r'^(?:'
    r'[IVXLC]+[\.\)]\s*'
    r'|[A-Za-z][\.\)]\s*'
    r'|\d+[\.\)]\s*'
    r')',
)

# Minor words for smart title-case (not capitalised unless first/last)
_MINOR_WORDS = frozenset({
    "a", "an", "the", "and", "but", "or", "for", "nor", "on",
    "at", "to", "by", "in", "of", "up", "as", "is", "if",
    "it", "vs", "via", "so", "yet",
})


# ══════════════════════════════════════════════════════════════
#  HeadingCounter — for IEEE-style numbered sections
# ══════════════════════════════════════════════════════════════

class _HeadingCounter:
    """Track heading numbers for IEEE-style section numbering (Roman/letter/arabic)."""

    _ROMAN = [
        "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
        "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX",
    ]
    _ALPHA = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")

    # Headings that should NOT receive automatic numbering
    _SKIP_LABELS = frozenset({
        "abstract", "references", "bibliography", "acknowledgments",
        "acknowledgements", "appendix", "appendices",
        "supplementary material", "supplementary materials",
        "supporting information",
    })

    def __init__(self):
        self.l1 = 0
        self.l2 = 0
        self.l3 = 0

    def should_number(self, heading_text: str) -> bool:
        text_lower = _EXISTING_NUM_RE.sub('', heading_text).strip().lower()
        return text_lower not in self._SKIP_LABELS

    def next_prefix(self, level: int) -> str:
        if level == 1:
            self.l1 += 1
            self.l2 = 0
            self.l3 = 0
            idx = self.l1 - 1
            roman = self._ROMAN[idx] if idx < len(self._ROMAN) else str(self.l1)
            return f"{roman}."
        elif level == 2:
            self.l2 += 1
            self.l3 = 0
            idx = self.l2 - 1
            alpha = self._ALPHA[idx] if idx < len(self._ALPHA) else str(self.l2)
            return f"{alpha}."
        elif level >= 3:
            self.l3 += 1
            return f"{self.l3})"
        return ""


# ══════════════════════════════════════════════════════════════
#  TransformerAgent
# ══════════════════════════════════════════════════════════════

class TransformerAgent:
    """Agent 5 — Apply formatting rules to produce a publication-ready DOCX."""

    # ── Public API ────────────────────────────────────────────

    def transform(
        self,
        docir: DocIR,
        style_spec: StyleSpec,
        input_path: Path,
        output_dir: Optional[Path] = None,
        formatted_bibliography: Optional[list[str]] = None,
    ) -> tuple[Path, list[ChangeRecord]]:
        """
        Apply *style_spec* formatting to the manuscript at *input_path*.

        Pipeline:
            0. Detect fragmentation → build merge plan
            1. Open original DOCX
            2. Execute paragraph merges in DOCX (combine text, remove extras)
            3. Build paragraph ↔ DocElement mapping
            4. Global settings (page layout, Normal style)
            5. Per-paragraph formatting based on detected role
            6. Structural inserts (labels, page breaks)
            7. Running head + page numbers
            8. Save formatted output

        Returns:
            ``(output_path, list_of_changes)``
        """
        if output_dir is None:
            output_dir = OUTPUT_FORMATTED_DIR
        output_dir.mkdir(parents=True, exist_ok=True)

        changes: list[ChangeRecord] = []

        # 0. Detect fragmentation & build merge plan ──────────
        merger = ParagraphMerger()
        merge_plan = merger.build_merge_plan(docir)

        # 1. Open ─────────────────────────────────────────────
        if input_path.suffix.lower() == ".txt":
            doc = Document()
            # For a brand new document from text, populate it with paragraphs
            # matching the DocIR elements so the rest of the pipeline works.
            for _ in docir.elements:
                doc.add_paragraph()
        else:
            doc = Document(str(input_path))

        typo = style_spec.default_typography
        layout = style_spec.page_layout

        # 2. Execute merges in DOCX ───────────────────────────
        if merge_plan:
            self._execute_merge_plan(doc, merge_plan)
            changes.append(ChangeRecord(
                category="paragraph_merge",
                description=f"Merged {sum(len(g) - 1 for g in merge_plan)} "
                            f"fragmented lines into {len(merge_plan)} paragraphs",
                rule_reference=f"{style_spec.style_name} paragraph reconstruction",
                status=ChangeStatus.APPLIED,
            ))

        # 2b. Clean PDF artifacts from run text ───────────────
        self._clean_run_artifacts(doc)

        # 3. Mapping (after merge — uses the merged DocIR) ───
        merged_docir = merger.apply_merge_plan(docir, merge_plan) if merge_plan else docir
        para_elem_map = self._build_para_elem_map_post_merge(merged_docir, merge_plan)

        # 4. Global settings ──────────────────────────────────
        changes.extend(self._apply_page_layout(doc, layout, style_spec))
        changes.extend(self._apply_default_typography(doc, typo, style_spec))

        # 5. Per-paragraph formatting ─────────────────────────
        body_count = 0
        ref_count = 0

        # Create heading counter for styles that use numbered sections
        heading_counter = (
            _HeadingCounter()
            if style_spec.headings.numbering_style == "ieee"
            else None
        )

        for idx, para in enumerate(doc.paragraphs):
            elem = para_elem_map.get(idx)
            if elem is None:
                # Paragraph not mapped — apply body defaults
                self._apply_body_defaults(para, style_spec)
                continue
            role = elem.role

            # Skip suppressed elements (page headers/footers cleared by structure detector)
            if role == ElementRole.UNKNOWN and not elem.content.strip():
                # Clear the DOCX paragraph text for suppressed elements
                for run in para.runs:
                    run.text = ""
                continue

            if role == ElementRole.TITLE:
                changes.extend(self._format_title(para, style_spec))
            elif role == ElementRole.AUTHOR_INFO:
                changes.extend(self._format_author_info(para, style_spec))
            elif role == ElementRole.ABSTRACT_LABEL:
                changes.extend(self._format_abstract_label(para, style_spec))
            elif role == ElementRole.ABSTRACT_BODY:
                changes.extend(self._format_abstract_body(para, style_spec))
            elif role == ElementRole.KEYWORDS:
                changes.extend(self._format_keywords(para, style_spec))
            elif role in _HEADING_ROLES:
                level = int(role.value.split("_")[1])
                changes.extend(self._format_heading(para, level, style_spec, heading_counter))
            elif role == ElementRole.REFERENCE_LABEL:
                changes.extend(self._format_reference_label(para, style_spec))
            elif role == ElementRole.REFERENCE_ENTRY:
                self._format_single_reference_entry(para, style_spec)
                # Replace text with citeproc-formatted version if available
                if formatted_bibliography and ref_count < len(formatted_bibliography):
                    fmt_text = formatted_bibliography[ref_count].strip()
                    if fmt_text:
                        self._replace_paragraph_text(para, fmt_text, style_spec)
                ref_count += 1
            elif role == ElementRole.BODY:
                self._apply_body_defaults(para, style_spec)
                body_count += 1
            elif role == ElementRole.UNKNOWN:
                # Treat remaining unknowns as body
                self._apply_body_defaults(para, style_spec)
            elif role == ElementRole.TABLE_CAPTION:
                changes.extend(self._format_table_caption(para, style_spec))
            elif role == ElementRole.FIGURE_CAPTION:
                changes.extend(self._format_figure_caption(para, style_spec))
            else:
                # Any other role (appendix…)
                self._apply_body_defaults(para, style_spec)

        if body_count:
            changes.append(ChangeRecord(
                category="body",
                description=f"{body_count} body paragraphs formatted: "
                            f"{typo.font_name} {typo.font_size_pt}pt, "
                            f"{typo.line_spacing}x spacing, "
                            f"{typo.first_line_indent_inches}\" indent",
                rule_reference=f"{style_spec.style_name} body typography",
                status=ChangeStatus.APPLIED,
            ))
        if ref_count:
            indent_desc = f"hanging indent {style_spec.references.hanging_indent_inches}\"" if style_spec.references.hanging_indent_inches > 0 else "no indent"
            changes.append(ChangeRecord(
                category="references",
                description=f"{ref_count} reference entries formatted: "
                            f"{indent_desc}, "
                            f"{style_spec.references.line_spacing}x spacing",
                rule_reference=f"{style_spec.style_name} reference list formatting",
                status=ChangeStatus.APPLIED,
            ))

        # 6. Structural inserts ───────────────────────────────
        title_text = self._get_title_text(merged_docir)
        changes.extend(
            self._add_structural_elements(doc, merged_docir, para_elem_map, style_spec)
        )

        # 7. Running head + page numbers ──────────────────────
        changes.extend(self._apply_running_head_and_page_numbers(doc, style_spec, title_text))

        # 7.5. Multi-column layout (e.g. IEEE two-column) ────
        changes.extend(self._apply_column_layout(doc, style_spec, para_elem_map))

        # 8. Save ─────────────────────────────────────────────
        stem = input_path.stem
        output_path = output_dir / f"{stem}_formatted.docx"
        doc.save(str(output_path))
        logger.info(
            "Formatted document saved → %s  (%d changes recorded)",
            output_path, len(changes),
        )
        return output_path, changes

    # ══════════════════════════════════════════════════════════
    #  Mapping helpers
    # ══════════════════════════════════════════════════════════

    @staticmethod
    def _build_para_elem_map(docir: DocIR) -> dict[int, DocElement]:
        """Map paragraph index in ``doc.paragraphs`` → DocElement."""
        mapping: dict[int, DocElement] = {}
        p_idx = 0
        for elem in docir.elements:
            if elem.type == ElementType.PARAGRAPH:
                mapping[p_idx] = elem
                p_idx += 1
            # Tables are not in doc.paragraphs — skip
        return mapping

    @staticmethod
    def _get_title_text(docir: DocIR) -> str:
        """Join all TITLE elements into a single string."""
        parts = [e.content for e in docir.elements if e.role == ElementRole.TITLE]
        return " ".join(parts).strip()

    # ── PDF artifact cleanup ──────────────────────────────────

    @staticmethod
    def _clean_run_artifacts(doc: Document) -> None:
        """Remove zero-width spaces, soft hyphens and other PDF artifacts from all runs.
        Also normalise leading/trailing whitespace so merged paragraphs are clean."""
        _ZWS = '\u200b'
        _RLM = '\u200f'
        _BOM = '\ufeff'
        _SHY = '\u00ad'
        for para in doc.paragraphs:
            for run in para.runs:
                t = run.text
                if _ZWS in t or _RLM in t or _BOM in t or _SHY in t:
                    t = t.replace(_ZWS, '').replace(_RLM, '')
                    t = t.replace(_BOM, '').replace(_SHY, '')
                    run.text = t
            # Strip leading whitespace from the first non-empty run
            for run in para.runs:
                if run.text:
                    stripped = run.text.lstrip()
                    if stripped != run.text:
                        run.text = stripped
                    break
            # Strip trailing whitespace from the last non-empty run
            for run in reversed(para.runs):
                if run.text:
                    stripped = run.text.rstrip()
                    if stripped != run.text:
                        run.text = stripped
                    break

    # ── Merge-plan execution ──────────────────────────────────

    def _execute_merge_plan(self, doc: Document, merge_plan: list[list[int]]) -> None:
        """
        Execute the merge plan on the actual DOCX paragraphs.

        For each group of paragraph indices, append subordinate text
        to the primary paragraph, then remove the subordinate
        paragraph elements from the XML tree.
        """
        if not merge_plan:
            return

        paras = doc.paragraphs          # snapshot of current paragraphs

        subordinate_indices: set[int] = set()

        for group in merge_plan:
            primary_idx = group[0]
            if primary_idx >= len(paras):
                continue
            primary_para = paras[primary_idx]

            for sub_idx in group[1:]:
                if sub_idx >= len(paras):
                    continue
                sub_para = paras[sub_idx]
                sub_text = sub_para.text.strip()
                subordinate_indices.add(sub_idx)
                if not sub_text:
                    continue

                # ── join logic (mirrors ParagraphMerger._join_texts) ──
                cur_text = primary_para.text
                if cur_text.endswith("-") and sub_text and sub_text[0].islower():
                    # hyphen-word reconstruction
                    if primary_para.runs:
                        last_run = primary_para.runs[-1]
                        if last_run.text.endswith("-"):
                            last_run.text = last_run.text[:-1]
                    primary_para.add_run(sub_text)
                else:
                    primary_para.add_run(" " + sub_text)

        # Remove subordinate paragraphs from the XML (reverse order)
        for idx in sorted(subordinate_indices, reverse=True):
            if idx >= len(paras):
                continue
            p_elem = paras[idx]._element
            parent = p_elem.getparent()
            if parent is not None:
                parent.remove(p_elem)

        logger.info(
            "Executed merge plan: removed %d subordinate paragraphs from DOCX",
            len(subordinate_indices),
        )

    def _build_para_elem_map_post_merge(
        self,
        merged_docir: DocIR,
        merge_plan: list[list[int]],
    ) -> dict[int, DocElement]:
        """
        Build paragraph → DocElement map that accounts for merged paragraphs.

        After ``_execute_merge_plan`` removes subordinate paragraphs from the
        DOCX, the remaining paragraphs correspond to *non-subordinate*
        DocIR elements. This method skips subordinate indices.
        """
        if not merge_plan:
            return self._build_para_elem_map(merged_docir)

        # Collect all subordinate paragraph indices
        subordinate_indices: set[int] = set()
        for group in merge_plan:
            for idx in group[1:]:
                subordinate_indices.add(idx)

        mapping: dict[int, DocElement] = {}
        docx_p_idx = 0          # post-merge DOCX paragraph counter
        docir_p_idx = 0         # original DocIR paragraph counter

        for elem in merged_docir.elements:
            if elem.type != ElementType.PARAGRAPH:
                continue
            if docir_p_idx in subordinate_indices:
                # This element was merged away — not in the DOCX anymore
                docir_p_idx += 1
                continue
            mapping[docx_p_idx] = elem
            docx_p_idx += 1
            docir_p_idx += 1

        return mapping

    # ══════════════════════════════════════════════════════════
    #  Global settings
    # ══════════════════════════════════════════════════════════

    def _apply_page_layout(self, doc: Document, layout, style_spec: StyleSpec) -> list[ChangeRecord]:
        """Set margins, page size, orientation on every section."""
        changes: list[ChangeRecord] = []
        for section in doc.sections:
            old_top = round(section.top_margin / 914400, 2) if section.top_margin else None
            section.top_margin = Inches(layout.margin_top_inches)
            section.bottom_margin = Inches(layout.margin_bottom_inches)
            section.left_margin = Inches(layout.margin_left_inches)
            section.right_margin = Inches(layout.margin_right_inches)
            section.page_width = Inches(layout.page_width_inches)
            section.page_height = Inches(layout.page_height_inches)
            changes.append(ChangeRecord(
                category="page_layout",
                description=(
                    f"Margins: {layout.margin_top_inches}\" T/B/L/R, "
                    f"page {layout.page_width_inches}\"×{layout.page_height_inches}\""
                ),
                old_value=f"top={old_top}\"" if old_top else "unknown",
                new_value=f"{layout.margin_top_inches}\" all sides",
                rule_reference=f"{style_spec.style_name} page layout",
                status=ChangeStatus.APPLIED,
            ))
        return changes

    def _apply_default_typography(self, doc: Document, typo, style_spec: StyleSpec) -> list[ChangeRecord]:
        """Set the Normal style defaults — every un-overridden paragraph inherits this."""
        changes: list[ChangeRecord] = []
        normal = doc.styles["Normal"]
        old_font = normal.font.name
        old_size = normal.font.size

        normal.font.name = typo.font_name
        normal.font.size = Pt(typo.font_size_pt)
        # Set CS / East-Asian font families so non-Latin text also uses the right font
        self._set_style_rfonts(normal, typo.font_name)

        pf = normal.paragraph_format
        pf.line_spacing = typo.line_spacing
        pf.first_line_indent = Inches(typo.first_line_indent_inches)
        pf.space_after = Pt(typo.space_after_paragraph_pt)
        pf.space_before = Pt(typo.space_before_paragraph_pt)

        changes.append(ChangeRecord(
            category="typography",
            description=f"Normal style → {typo.font_name} {typo.font_size_pt}pt, "
                        f"{typo.line_spacing}× spacing",
            old_value=f"{old_font} {old_size}",
            new_value=f"{typo.font_name} {typo.font_size_pt}pt",
            rule_reference=f"{style_spec.style_name} typography",
            status=ChangeStatus.APPLIED,
        ))
        return changes

    # ══════════════════════════════════════════════════════════
    #  Font enforcement helpers
    # ══════════════════════════════════════════════════════════

    @staticmethod
    def _set_style_rfonts(style, font_name: str) -> None:
        """Set all four rFonts families on a *style* element."""
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), font_name)

    @staticmethod
    def _enforce_run_font(run, font_name: str, font_size_pt: float) -> None:
        """Force font name + size on a single run (overrides inherited junk)."""
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        # Also set cs / eastAsia via XML for full coverage
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), font_name)

    def _enforce_font(
        self,
        para,
        font_name: str,
        font_size_pt: float,
        *,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
    ) -> None:
        """Force font on **every** run in *para*, optionally setting bold/italic."""
        for run in para.runs:
            self._enforce_run_font(run, font_name, font_size_pt)
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic

    # ══════════════════════════════════════════════════════════
    #  Text-case + paragraph-text helpers
    # ══════════════════════════════════════════════════════════

    @staticmethod
    def _apply_text_case(text: str, case: str) -> str:
        """Transform heading text to the specified case style."""
        if not text:
            return text
        if case == "upper":
            return text.upper()
        elif case == "title_case":
            words = text.split()
            result: list[str] = []
            for i, w in enumerate(words):
                if i == 0 or i == len(words) - 1 or w.lower() not in _MINOR_WORDS:
                    # Preserve ALL-CAPS abbreviations (DNA, RNA, EHEC, etc.)
                    if w.isupper() and len(w) > 1:
                        result.append(w)  # keep abbreviations as-is
                    else:
                        result.append(w.capitalize())
                else:
                    result.append(w.lower())
            return " ".join(result)
        elif case == "sentence_case":
            if len(text) <= 1:
                return text.upper()
            # Preserve ALL-CAPS abbreviations within the sentence
            words = text.split()
            result = []
            for i, w in enumerate(words):
                if i == 0:
                    result.append(w.capitalize())
                elif w.isupper() and len(w) > 1:
                    result.append(w)  # preserve abbreviations
                else:
                    result.append(w.lower())
            return " ".join(result)
        return text

    def _set_paragraph_text(
        self, para, new_text: str,
        font_name: str, font_size_pt: float,
        *, bold: Optional[bool] = None, italic: Optional[bool] = None,
    ) -> None:
        """Replace all text in a paragraph with *new_text* and enforce font."""
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
            self._enforce_run_font(para.runs[0], font_name, font_size_pt)
            if bold is not None:
                para.runs[0].bold = bold
            if italic is not None:
                para.runs[0].italic = italic
        else:
            run = para.add_run(new_text)
            self._enforce_run_font(run, font_name, font_size_pt)
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic

    # ── Title ─────────────────────────────────────────────────

    def _format_title(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        ts = spec.title_page.title
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = _ALIGN.get(ts.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = spec.default_typography.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        self._enforce_font(
            para,
            spec.default_typography.font_name,
            ts.font_size_pt,
            bold=ts.bold,
            italic=ts.italic,
        )
        return [ChangeRecord(
            category="title_page",
            description=f"Title: {ts.alignment}, {'bold' if ts.bold else 'plain'}, {ts.font_size_pt}pt",
            rule_reference=f"{spec.style_name} title formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ── Author info ───────────────────────────────────────────

    def _format_author_info(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        an = spec.title_page.author_name
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = _ALIGN.get(an.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = spec.default_typography.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        self._enforce_font(
            para,
            spec.default_typography.font_name,
            an.font_size_pt,
            bold=an.bold,
            italic=an.italic,
        )
        return [ChangeRecord(
            category="title_page",
            description=f"Author info: {an.alignment}, {an.font_size_pt}pt",
            rule_reference=f"{spec.style_name} author formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ── Abstract label ────────────────────────────────────────

    def _format_abstract_label(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = _ALIGN.get(spec.abstract.label_alignment, WD_ALIGN_PARAGRAPH.CENTER)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = spec.default_typography.line_spacing
        pf.space_after = Pt(0)
        self._enforce_font(
            para,
            spec.default_typography.font_name,
            spec.default_typography.font_size_pt,
            bold=spec.abstract.label_bold,
        )
        return [ChangeRecord(
            category="abstract",
            description=f"Abstract label: {spec.abstract.label_alignment}, {'bold' if spec.abstract.label_bold else 'plain'}",
            rule_reference=f"{spec.style_name} abstract formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ── Abstract body ─────────────────────────────────────────

    def _format_abstract_body(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        # Abstract body: use body alignment but NO first-line indent
        # unless the style spec explicitly requires it
        para.alignment = _ALIGN.get(spec.default_typography.paragraph_alignment, WD_ALIGN_PARAGRAPH.LEFT)
        pf = para.paragraph_format
        if spec.abstract.paragraph_indent:
            pf.first_line_indent = Inches(spec.default_typography.first_line_indent_inches)
        else:
            # Explicitly set to 0 via XML to override inherited indent
            pf.first_line_indent = Pt(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = spec.default_typography.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        self._enforce_font(
            para,
            spec.default_typography.font_name,
            spec.default_typography.font_size_pt,
            bold=False,
        )
        return [ChangeRecord(
            category="abstract",
            description=f"Abstract body: {spec.default_typography.paragraph_alignment}, no indent, {spec.default_typography.line_spacing}x spacing",
            rule_reference=f"{spec.style_name} abstract formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ── Keywords ──────────────────────────────────────────────

    def _format_keywords(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = _ALIGN.get(spec.default_typography.paragraph_alignment, WD_ALIGN_PARAGRAPH.LEFT)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(0.5) if spec.abstract.keywords_indent else Inches(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = spec.default_typography.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        # Apply italic to "Keywords:" prefix, plain to the rest
        for run in para.runs:
            self._enforce_run_font(
                run,
                spec.default_typography.font_name,
                spec.default_typography.font_size_pt,
            )
            if "keyword" in run.text.lower():
                run.bold = False
                run.italic = True
            else:
                run.bold = False
                run.italic = False
        return [ChangeRecord(
            category="abstract",
            description=f"Keywords: {'indented 0.5\"' if spec.abstract.keywords_indent else 'flush left'}, {'italic prefix' if spec.abstract.keywords_italic else 'plain prefix'}",
            rule_reference=f"{spec.style_name} keywords formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ── Headings (all 5 levels) ───────────────────────────────

    def _format_heading(
        self, para, level: int, spec: StyleSpec,
        heading_counter: Optional[_HeadingCounter] = None,
    ) -> list[ChangeRecord]:
        hs = spec.headings.get_level(level)
        typo = spec.default_typography
        pf = para.paragraph_format

        # Alignment
        if hs.alignment == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Inches(0)
            pf.left_indent = Inches(0)
        elif hs.alignment == "left":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Inches(0)
            pf.left_indent = Inches(0)
        elif hs.alignment == "indented":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Inches(hs.indent_inches or 0.5)
            pf.left_indent = Inches(0)

        pf.line_spacing = typo.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        # ── Text case transformation ─────────────────────────
        original_text = para.text.strip()
        clean_text = _EXISTING_NUM_RE.sub('', original_text).strip()

        if hs.case:
            new_text = self._apply_text_case(clean_text, hs.case)
        else:
            new_text = clean_text

        # ── Heading numbering (IEEE style) ───────────────────
        if heading_counter and heading_counter.should_number(clean_text):
            prefix = heading_counter.next_prefix(level)
            new_text = f"{prefix} {new_text}"

        # Replace paragraph text if it changed
        if new_text != original_text:
            self._set_paragraph_text(
                para, new_text,
                typo.font_name, hs.font_size_pt,
                bold=hs.bold, italic=hs.italic,
            )
        else:
            self._enforce_font(
                para,
                typo.font_name,
                hs.font_size_pt,
                bold=hs.bold,
                italic=hs.italic,
            )

        return [ChangeRecord(
            category="headings",
            description=f"Level {level}: {hs.description}"
                        + (f" [{new_text[:40]}]" if new_text != original_text else ""),
            rule_reference=f"{spec.style_name} heading level {level}",
            status=ChangeStatus.APPLIED,
        )]

    # ── Reference label ───────────────────────────────────────

    def _format_reference_label(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        """Format the 'References' / 'Works Cited' / 'Bibliography' label per style spec."""
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        # Replace the label text with the style-specific label
        expected_label = spec.references.section_label
        current_text = para.text.strip()
        if current_text.lower() != expected_label.lower():
            # Clear existing runs and set new text
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = expected_label
            else:
                para.add_run(expected_label)

        para.alignment = _ALIGN.get(spec.references.label_alignment, WD_ALIGN_PARAGRAPH.CENTER)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = spec.default_typography.line_spacing
        pf.space_after = Pt(0)
        self._enforce_font(
            para,
            spec.default_typography.font_name,
            spec.default_typography.font_size_pt,
            bold=spec.references.label_bold,
        )
        return [ChangeRecord(
            category="references",
            description=f"\"{spec.references.section_label}\" label: {spec.references.label_alignment}, {'bold' if spec.references.label_bold else 'plain'}",
            rule_reference=f"{spec.style_name} reference list formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ── Single reference entry ────────────────────────────────

    def _format_single_reference_entry(self, para, spec: StyleSpec) -> None:
        """Format one reference entry per style spec (hanging/no indent, spacing)."""
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = _ALIGN.get(spec.default_typography.paragraph_alignment, WD_ALIGN_PARAGRAPH.LEFT)
        pf = para.paragraph_format
        if spec.references.entry_indent_type == "hanging" and spec.references.hanging_indent_inches > 0:
            pf.left_indent = Inches(spec.references.hanging_indent_inches)
            pf.first_line_indent = Inches(-spec.references.hanging_indent_inches)
        else:
            pf.left_indent = Inches(0)
            pf.first_line_indent = Inches(0)
        pf.line_spacing = spec.references.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        self._enforce_font(
            para,
            spec.default_typography.font_name,
            spec.default_typography.font_size_pt,
            bold=False,
        )

    # ── Replace paragraph text with formatted reference ────────

    def _replace_paragraph_text(
        self, para, new_text: str, spec: StyleSpec,
    ) -> None:
        """Replace all runs in a paragraph with a single run of *new_text*."""
        # Clear existing runs
        for run in para.runs:
            run.text = ""
        # If there are runs, use the first one; otherwise add a new run
        if para.runs:
            para.runs[0].text = new_text
            self._enforce_run_font(
                para.runs[0],
                spec.default_typography.font_name,
                spec.default_typography.font_size_pt,
            )
            para.runs[0].bold = False
        else:
            run = para.add_run(new_text)
            self._enforce_run_font(
                run,
                spec.default_typography.font_name,
                spec.default_typography.font_size_pt,
            )
            run.bold = False

    # ── Body defaults ─────────────────────────────────────────

    def _apply_body_defaults(self, para, spec: StyleSpec) -> None:
        """Apply standard body-paragraph formatting (no ChangeRecord)."""
        typo = spec.default_typography
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = _ALIGN.get(typo.paragraph_alignment, WD_ALIGN_PARAGRAPH.LEFT)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(typo.first_line_indent_inches)
        pf.left_indent = Inches(0)
        pf.line_spacing = typo.line_spacing
        pf.space_after = Pt(typo.space_after_paragraph_pt)
        pf.space_before = Pt(typo.space_before_paragraph_pt)
        self._enforce_font(
            para,
            typo.font_name,
            typo.font_size_pt,
        )

    # ── Table caption ──────────────────────────────────────────

    def _format_table_caption(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        """Format a table caption per style spec (e.g. 'Table 1' bold, title italic)."""
        typo = spec.default_typography
        tbl = spec.tables
        align_key = typo.paragraph_alignment.lower()
        para.alignment = _ALIGN.get(align_key, WD_ALIGN_PARAGRAPH.LEFT)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = typo.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        # Try to style the "Table N" part bold and the title italic
        text = para.text
        for run in para.runs:
            self._enforce_run_font(run, typo.font_name, typo.font_size_pt)
            run_lower = run.text.strip().lower()
            if run_lower.startswith(tbl.number_label.lower()):
                run.bold = tbl.number_bold
                run.italic = tbl.number_italic
            else:
                run.bold = False
                run.italic = tbl.title_italic

        return [ChangeRecord(
            category="tables_figures",
            description=f"Table caption formatted: {tbl.number_label} {'bold' if tbl.number_bold else 'plain'}, title {'italic' if tbl.title_italic else 'plain'}",
            rule_reference=f"{spec.style_name} table formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ── Figure caption ─────────────────────────────────────────

    def _format_figure_caption(self, para, spec: StyleSpec) -> list[ChangeRecord]:
        """Format a figure caption per style spec."""
        typo = spec.default_typography
        fig = spec.figures
        align_key = typo.paragraph_alignment.lower()
        para.alignment = _ALIGN.get(align_key, WD_ALIGN_PARAGRAPH.LEFT)
        pf = para.paragraph_format
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0)
        pf.line_spacing = typo.line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        for run in para.runs:
            self._enforce_run_font(run, typo.font_name, typo.font_size_pt)
            run_lower = run.text.strip().lower()
            if run_lower.startswith(fig.number_label.lower()):
                run.bold = fig.number_bold
                run.italic = fig.number_italic
            else:
                run.bold = False
                run.italic = fig.title_italic

        return [ChangeRecord(
            category="tables_figures",
            description=f"Figure caption formatted: {fig.number_label} {'bold' if fig.number_bold else 'plain'}, title {'italic' if fig.title_italic else 'plain'}",
            rule_reference=f"{spec.style_name} figure formatting",
            status=ChangeStatus.APPLIED,
        )]

    # ══════════════════════════════════════════════════════════
    #  Structural inserts (labels, page breaks)
    # ══════════════════════════════════════════════════════════

    def _add_structural_elements(
        self,
        doc: Document,
        docir: DocIR,
        para_elem_map: dict[int, DocElement],
        spec: StyleSpec,
    ) -> list[ChangeRecord]:
        """Insert missing labels and APA-required page breaks."""
        changes: list[ChangeRecord] = []

        has_abstract_label = any(
            e.role == ElementRole.ABSTRACT_LABEL for e in docir.elements
        )
        has_reference_label = any(
            e.role == ElementRole.REFERENCE_LABEL for e in docir.elements
        )

        # Collect role spans ──────────────────────────────────
        first_abstract_idx: Optional[int] = None
        first_ref_idx: Optional[int] = None
        first_body_or_heading_after_abstract: Optional[int] = None
        last_front_matter_idx: Optional[int] = None  # last title/author

        abstract_seen = False
        for idx, elem in para_elem_map.items():
            role = elem.role
            if role in (ElementRole.TITLE, ElementRole.AUTHOR_INFO):
                last_front_matter_idx = idx
            if role == ElementRole.ABSTRACT_BODY and first_abstract_idx is None:
                first_abstract_idx = idx
                abstract_seen = True
            if role == ElementRole.ABSTRACT_LABEL and first_abstract_idx is None:
                first_abstract_idx = idx
                abstract_seen = True
            if role == ElementRole.REFERENCE_ENTRY and first_ref_idx is None:
                first_ref_idx = idx
            if abstract_seen and role in (
                ElementRole.BODY,
                ElementRole.HEADING_1, ElementRole.HEADING_2,
                ElementRole.HEADING_3,
            ):
                if first_body_or_heading_after_abstract is None:
                    first_body_or_heading_after_abstract = idx

        paras = doc.paragraphs

        # ── Insert "Abstract" label if missing ───────────────
        _ALIGN = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        # Page break before abstract only when style requires a title page
        needs_section_breaks = spec.title_page.required

        if not has_abstract_label and first_abstract_idx is not None:
            target = paras[first_abstract_idx]
            new_para = self._insert_paragraph_before(target, doc)
            run = new_para.add_run(spec.abstract.label)
            self._enforce_run_font(
                run,
                spec.default_typography.font_name,
                spec.default_typography.font_size_pt,
            )
            run.bold = spec.abstract.label_bold
            new_para.alignment = _ALIGN.get(spec.abstract.label_alignment, WD_ALIGN_PARAGRAPH.CENTER)
            new_para.paragraph_format.first_line_indent = Inches(0)
            new_para.paragraph_format.left_indent = Inches(0)
            new_para.paragraph_format.line_spacing = spec.default_typography.line_spacing
            new_para.paragraph_format.space_after = Pt(0)
            new_para.paragraph_format.space_before = Pt(0)
            if needs_section_breaks:
                new_para.paragraph_format.page_break_before = True
            changes.append(ChangeRecord(
                category="abstract",
                description=f"Inserted \"{spec.abstract.label}\" label ({spec.abstract.label_alignment}, {'bold' if spec.abstract.label_bold else 'plain'})"
                            + (" on new page" if needs_section_breaks else ""),
                rule_reference=f"{spec.style_name} abstract formatting",
                status=ChangeStatus.APPLIED,
            ))
        elif first_abstract_idx is not None:
            if needs_section_breaks:
                paras[first_abstract_idx].paragraph_format.page_break_before = True

        # ── Insert "References" label if missing ─────────────
        if not has_reference_label and first_ref_idx is not None:
            target = paras[first_ref_idx]
            new_para = self._insert_paragraph_before(target, doc)
            run = new_para.add_run(spec.references.section_label)
            self._enforce_run_font(
                run,
                spec.default_typography.font_name,
                spec.default_typography.font_size_pt,
            )
            run.bold = spec.references.label_bold
            new_para.alignment = _ALIGN.get(spec.references.label_alignment, WD_ALIGN_PARAGRAPH.CENTER)
            new_para.paragraph_format.first_line_indent = Inches(0)
            new_para.paragraph_format.left_indent = Inches(0)
            new_para.paragraph_format.line_spacing = spec.default_typography.line_spacing
            new_para.paragraph_format.space_after = Pt(0)
            # References on new page
            new_para.paragraph_format.page_break_before = True
            changes.append(ChangeRecord(
                category="references",
                description=f"Inserted \"{spec.references.section_label}\" label ({spec.references.label_alignment}, {'bold' if spec.references.label_bold else 'plain'}) on new page",
                rule_reference=f"{spec.style_name} reference list formatting",
                status=ChangeStatus.APPLIED,
            ))
        elif first_ref_idx is not None:
            paras[first_ref_idx].paragraph_format.page_break_before = True

        # ── Title-page break ─────────────────────────────────
        #    Body text starts on a new page after abstract/keywords
        #    Only when the style requires separate sections (title page)
        if needs_section_breaks and first_body_or_heading_after_abstract is not None:
            paras[first_body_or_heading_after_abstract].paragraph_format.page_break_before = True
            changes.append(ChangeRecord(
                category="page_layout",
                description="Body text starts on new page (after abstract)",
                rule_reference=f"{spec.style_name} section breaks",
                status=ChangeStatus.APPLIED,
            ))

        return changes

    # ── XML paragraph insertion ───────────────────────────────

    @staticmethod
    def _insert_paragraph_before(target_para, doc: Document) -> Paragraph:
        """Insert a new empty paragraph immediately before *target_para*."""
        new_p = OxmlElement("w:p")
        target_para._element.addprevious(new_p)
        return Paragraph(new_p, target_para._element.getparent())

    # ══════════════════════════════════════════════════════════
    #  Running head + page numbers
    # ══════════════════════════════════════════════════════════

    def _apply_running_head_and_page_numbers(
        self,
        doc: Document,
        spec: StyleSpec,
        title_text: str,
    ) -> list[ChangeRecord]:
        """
        Add page numbers (always) and running-head text (only if enabled).

        APA 7 student papers: page number only (top-right).
        APA 7 professional papers: shortened title (left) + page number (right).
        """
        changes: list[ChangeRecord] = []
        font_name = spec.default_typography.font_name
        font_size = spec.running_head.font_size_pt
        running_head_enabled = spec.running_head.enabled

        for section in doc.sections:
            section.different_first_page_header_footer = False
            header = section.header
            header.is_linked_to_previous = False

            # Clear any existing content
            for p in header.paragraphs:
                p_elem = p._element
                for child in list(p_elem):
                    p_elem.remove(child)

            if not header.paragraphs:
                header.add_paragraph()

            h_para = header.paragraphs[0]

            # ── Tab stop at right margin ─────────────────────
            usable_width_twips = int(
                (
                    spec.page_layout.page_width_inches
                    - spec.page_layout.margin_left_inches
                    - spec.page_layout.margin_right_inches
                )
                * 1440  # 1 inch = 1440 twips
            )
            pPr = h_para._element.get_or_add_pPr()
            tabs_elem = OxmlElement("w:tabs")
            tab_elem = OxmlElement("w:tab")
            tab_elem.set(qn("w:val"), "right")
            tab_elem.set(qn("w:pos"), str(usable_width_twips))
            tab_elem.set(qn("w:leader"), "none")
            tabs_elem.append(tab_elem)
            pPr.append(tabs_elem)

            # ── Running head text (left) — only if enabled ───
            short_title = ""
            if running_head_enabled:
                short_title = title_text.upper()[:spec.running_head.max_characters]
                if short_title:
                    head_run = h_para.add_run(short_title)
                    self._enforce_run_font(head_run, font_name, font_size)

            # ── Tab ──────────────────────────────────────────
            tab_run = h_para.add_run("\t")
            self._enforce_run_font(tab_run, font_name, font_size)

            # ── Page number field (always) ───────────────────
            self._add_page_number_field(h_para, font_name, font_size)

            if running_head_enabled:
                changes.append(ChangeRecord(
                    category="running_head",
                    description=f"Running head: \"{short_title}\" + page number",
                    rule_reference=f"{spec.style_name} running head",
                    status=ChangeStatus.APPLIED,
                ))
            else:
                changes.append(ChangeRecord(
                    category="page_layout",
                    description="Page number added (top-right, every page)",
                    rule_reference=f"{spec.style_name} page numbering",
                    status=ChangeStatus.APPLIED,
                ))

        return changes

    # ──────────────────────────────────────────────────────────
    #  Multi-column layout (IEEE two-column, etc.)
    # ──────────────────────────────────────────────────────────

    def _apply_column_layout(
        self,
        doc,
        spec: "StyleSpec",
        para_elem_map: dict[int, "DocElement"],
    ) -> list["ChangeRecord"]:
        """Apply multi-column layout.

        For IEEE: single-column title / abstract → continuous section break
        → two-column body.  For other styles (columns == 1) this is a no-op.
        """
        num_cols = getattr(spec.page_layout, "columns", 1)
        if num_cols <= 1:
            return []

        col_gap_twips = int(
            getattr(spec.page_layout, "column_spacing_inches", 0.25) * 1440
        )
        layout = spec.page_layout
        changes: list[ChangeRecord] = []

        # --- locate last "front-matter" paragraph (title / author / abstract / kw)
        _FRONT_ROLES = frozenset({
            ElementRole.TITLE,
            ElementRole.AUTHOR_INFO,
            ElementRole.ABSTRACT_LABEL,
            ElementRole.ABSTRACT_BODY,
            ElementRole.KEYWORDS,
        })
        last_front_idx: int | None = None
        for idx in sorted(para_elem_map.keys()):
            elem = para_elem_map[idx]
            if elem.role in _FRONT_ROLES:
                last_front_idx = idx

        paras = doc.paragraphs

        if last_front_idx is not None and last_front_idx < len(paras):
            # ── Insert continuous section break after front-matter ──
            front_para = paras[last_front_idx]
            pPr = front_para._element.get_or_add_pPr()

            # Remove any existing sectPr from this paragraph
            for old in pPr.findall(qn("w:sectPr")):
                pPr.remove(old)

            first_sectPr = OxmlElement("w:sectPr")
            # continuous break → no page break, just column change
            sect_type = OxmlElement("w:type")
            sect_type.set(qn("w:val"), "continuous")
            first_sectPr.append(sect_type)
            # first section stays single-column
            first_cols = OxmlElement("w:cols")
            first_cols.set(qn("w:num"), "1")
            first_sectPr.append(first_cols)
            # carry page size
            pgSz = OxmlElement("w:pgSz")
            pgSz.set(qn("w:w"), str(int(layout.page_width_inches * 1440)))
            pgSz.set(qn("w:h"), str(int(layout.page_height_inches * 1440)))
            first_sectPr.append(pgSz)
            # carry margins
            pgMar = OxmlElement("w:pgMar")
            pgMar.set(qn("w:top"), str(int(layout.margin_top_inches * 1440)))
            pgMar.set(qn("w:bottom"), str(int(layout.margin_bottom_inches * 1440)))
            pgMar.set(qn("w:left"), str(int(layout.margin_left_inches * 1440)))
            pgMar.set(qn("w:right"), str(int(layout.margin_right_inches * 1440)))
            first_sectPr.append(pgMar)

            pPr.append(first_sectPr)
            logger.info("Inserted continuous section break after paragraph %d", last_front_idx)

        # ── Set the FINAL section (body) to N columns ──────────
        final_sectPr = doc.sections[-1]._sectPr
        old_cols = final_sectPr.find(qn("w:cols"))
        if old_cols is not None:
            final_sectPr.remove(old_cols)
        body_cols = OxmlElement("w:cols")
        body_cols.set(qn("w:num"), str(num_cols))
        body_cols.set(qn("w:space"), str(col_gap_twips))
        # equal-width columns
        body_cols.set(qn("w:equalWidth"), "1")
        final_sectPr.append(body_cols)

        changes.append(ChangeRecord(
            category="page_layout",
            description=(
                f"{num_cols}-column layout applied "
                f"(single-column front matter → {num_cols}-column body)"
            ),
            rule_reference=f"{spec.style_name} column layout",
            status=ChangeStatus.APPLIED,
        ))
        return changes

    def _add_page_number_field(
        self, paragraph, font_name: str, font_size_pt: float,
    ) -> None:
        """Insert a PAGE field code (renders as the current page number)."""
        # Begin
        run_begin = paragraph.add_run()
        self._enforce_run_font(run_begin, font_name, font_size_pt)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._element.append(fld_begin)

        # InstrText
        run_instr = paragraph.add_run()
        self._enforce_run_font(run_instr, font_name, font_size_pt)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._element.append(instr)

        # End
        run_end = paragraph.add_run()
        self._enforce_run_font(run_end, font_name, font_size_pt)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._element.append(fld_end)


# ── Module-level constants ────────────────────────────────────

_HEADING_ROLES = frozenset({
    ElementRole.HEADING_1,
    ElementRole.HEADING_2,
    ElementRole.HEADING_3,
    ElementRole.HEADING_4,
    ElementRole.HEADING_5,
})
