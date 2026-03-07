"""
Document Editor Agent — Unified REST API + MCP Server
======================================================

Single-file FastAPI backend with embedded Agno agent + Groq LLM.
All document tools, MCP server, and REST endpoints in one place.

Run:
    uvicorn api:app --host 0.0.0.0 --port 8080 --reload
"""

from __future__ import annotations

import asyncio
import html as html_mod
import io
import json
import logging
import os
import re
import shutil
import zipfile
import time
import uuid
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from pydantic import BaseModel

# ── Logging ──────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO)
log = logging.getLogger("api")

# ── Load .env ────────────────────────────────────────────────────────────
_ENV_FILE = Path(__file__).resolve().parent / ".env"
if _ENV_FILE.exists():
    for line in _ENV_FILE.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            key, _, value = line.partition("=")
            os.environ.setdefault(key.strip(), value.strip())

# ── Paths ────────────────────────────────────────────────────────────────
# Use a directory outside the 'Assistant' folder to avoid uvicorn reload loops
DOCUMENTS_DIR = Path(r"D:\CODEDB\Hackathon\HackaMineD\HackaMineD\AgentCode\documents_store").resolve()
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
os.environ["DOCUMENTS_DIR"] = str(DOCUMENTS_DIR)


# ═════════════════════════════════════════════════════════════════════════
#  DOCX HELPERS  (must be defined BEFORE tools that use them)
# ═════════════════════════════════════════════════════════════════════════

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HIGHLIGHT_MAP = {
    "yellow": WD_COLOR_INDEX.YELLOW,
    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "pink": WD_COLOR_INDEX.PINK,
    "blue": WD_COLOR_INDEX.BLUE,
    "red": WD_COLOR_INDEX.RED,
    "turquoise": WD_COLOR_INDEX.TURQUOISE,
    "gray": WD_COLOR_INDEX.GRAY_25,
}


def _resolve(filename: str) -> Path:
    return DOCUMENTS_DIR / Path(filename).name


def _load(filename: str) -> Document:
    path = _resolve(filename)
    if not path.exists():
        raise FileNotFoundError(f"Document '{filename}' not found.")
    return Document(str(path))


def _save(doc: Document, filename: str) -> Path:
    path = _resolve(filename)
    doc.save(str(path))
    return path


def _apply_run_format(run, *, bold=None, italic=None, underline=None,
                      strikethrough=None, font_size=None, font_name=None,
                      font_color=None, highlight_color=None):
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    if strikethrough is not None:
        run.font.strike = strikethrough
    if font_size is not None:
        run.font.size = Pt(font_size)
    if font_name is not None:
        run.font.name = font_name
    if font_color is not None:
        run.font.color.rgb = RGBColor.from_string(font_color.lstrip("#"))
    if highlight_color is not None:
        hl = HIGHLIGHT_MAP.get(highlight_color.lower())
        if hl:
            run.font.highlight_color = hl


def _set_run_text(r_elem, text: str):
    for t in r_elem.findall(qn("w:t")):
        r_elem.remove(t)
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t_elem.set(qn("xml:space"), "preserve")
    r_elem.append(t_elem)


def _format_matching_text(paragraph, search_text: str, **fmt) -> bool:
    full_text = paragraph.text
    idx = full_text.find(search_text)
    if idx == -1:
        return False
    target_end = idx + len(search_text)
    p_elem = paragraph._p
    runs = list(paragraph.runs)
    boundaries: list[tuple] = []
    pos = 0
    for run in runs:
        rlen = len(run.text)
        boundaries.append((pos, pos + rlen, run))
        pos += rlen
    new_elements: list[tuple] = []
    for run_start, run_end, run in boundaries:
        overlap_start = max(run_start, idx)
        overlap_end = min(run_end, target_end)
        if overlap_start >= overlap_end:
            new_elements.append(("keep", run._r))
            continue
        if run_start < overlap_start:
            before_r = deepcopy(run._r)
            _set_run_text(before_r, run.text[: overlap_start - run_start])
            new_elements.append(("keep", before_r))
        target_r = deepcopy(run._r)
        _set_run_text(target_r, run.text[overlap_start - run_start: overlap_end - run_start])
        new_elements.append(("format", target_r))
        if overlap_end < run_end:
            after_r = deepcopy(run._r)
            _set_run_text(after_r, run.text[overlap_end - run_start:])
            new_elements.append(("keep", after_r))
    for run in runs:
        p_elem.remove(run._r)
    for action, r_elem in new_elements:
        p_elem.append(r_elem)
        if action == "format":
            from docx.text.run import Run
            temp_run = Run(r_elem, paragraph)
            _apply_run_format(temp_run, **fmt)
    return True


def _run_to_html(run) -> str:
    txt = html_mod.escape(run.text)
    if not txt:
        return ""
    styles: list[str] = []
    if run.bold:
        txt = f"<b>{txt}</b>"
    if run.italic:
        txt = f"<i>{txt}</i>"
    if run.underline:
        styles.append("text-decoration:underline")
    if run.font and run.font.strike:
        styles.append("text-decoration:line-through")
    if run.font and run.font.size:
        styles.append(f"font-size:{run.font.size.pt}pt")
    if run.font and run.font.name:
        styles.append(f"font-family:'{run.font.name}'")
    if run.font and run.font.color and run.font.color.rgb:
        styles.append(f"color:#{run.font.color.rgb}")
    if styles:
        txt = f'<span style="{";".join(styles)}">{txt}</span>'
    return txt


def _para_to_html(para, idx: int) -> str:
    name = (para.style.name if para.style else "Normal").lower()
    align = ""
    if para.alignment is not None:
        align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
        align = f"text-align:{align_map.get(para.alignment, 'left')};"
    tag_map = {
        "heading 1": ("h2", "margin:0.4em 0 0.2em;color:#1a73e8;"),
        "heading 2": ("h3", "margin:0.3em 0 0.15em;color:#1a73e8;"),
        "heading 3": ("h4", "margin:0.25em 0 0.1em;color:#555;"),
        "title": ("h1", "margin:0.5em 0 0.3em;color:#202124;"),
    }
    tag, base_style = "p", "margin:0.25em 0;line-height:1.6;"
    for key, (t, s) in tag_map.items():
        if key in name:
            tag, base_style = t, s
            break
    if "list" in name or "bullet" in name:
        tag, base_style = "li", "margin-left:1.5em;"
    inner = "".join(_run_to_html(r) for r in para.runs)
    if not inner:
        inner = html_mod.escape(para.text) if para.text else "&nbsp;"
    return f'<{tag} style="{base_style}{align}" data-para-index="{idx}">{inner}</{tag}>'


def _docx_to_html(doc: Document) -> str:
    parts = [
        '<div style="font-family:\'Segoe UI\',Calibri,sans-serif;'
        'font-size:11pt;color:#202124;max-width:100%;padding:0.5em;">'
    ]
    for i, para in enumerate(doc.paragraphs):
        parts.append(_para_to_html(para, i))
    for table in doc.tables:
        parts.append(
            '<table style="border-collapse:collapse;width:100%;margin:0.5em 0;font-size:10pt;">'
        )
        for r_idx, row in enumerate(table.rows):
            parts.append("<tr>")
            for cell in row.cells:
                tag_name = "th" if r_idx == 0 else "td"
                style = (
                    "border:1px solid #ddd;padding:6px 10px;"
                    + ("background:#f1f3f4;font-weight:600;" if r_idx == 0 else "")
                )
                parts.append(f'<{tag_name} style="{style}">{html_mod.escape(cell.text)}</{tag_name}>')
            parts.append("</tr>")
        parts.append("</table>")
    parts.append("</div>")
    return "\n".join(parts)


# ═════════════════════════════════════════════════════════════════════════
#  AGENT TOOL FUNCTIONS  (plain functions for Agno)
# ═════════════════════════════════════════════════════════════════════════

def list_documents(dummy: str = "list") -> str:
    """List all .docx documents in the store. Returns JSON with filenames and sizes.
    Always pass dummy='list' when calling this tool."""
    docs = []
    for f in sorted(DOCUMENTS_DIR.glob("*.docx")):
        stat = f.stat()
        docs.append({"filename": f.name, "size_bytes": stat.st_size, "modified": stat.st_mtime})
    return json.dumps(docs, indent=2)


def create_document(title: str) -> str:
    """Create a new blank .docx document with the given title as filename and first heading."""
    safe_title = re.sub(r'[<>:"/\\|?*]', "_", title)
    filename = f"{safe_title}.docx"
    doc = Document()
    doc.add_heading(title, level=1)
    _save(doc, filename)
    return json.dumps({"filename": filename, "message": f"Document '{title}' created."})


def delete_document(filename: str) -> str:
    """Permanently delete a document from the store."""
    path = _resolve(filename)
    if not path.exists():
        return json.dumps({"error": f"Document '{filename}' not found."})
    path.unlink()
    return json.dumps({"message": f"Document '{filename}' deleted."})


def get_document_info(filename: str) -> str:
    """Get metadata: paragraph count, word count, styles used, table count."""
    doc = _load(filename)
    paras = doc.paragraphs
    word_count = sum(len(p.text.split()) for p in paras)
    styles_used = list({p.style.name for p in paras if p.style})
    return json.dumps({
        "filename": filename,
        "paragraph_count": len(paras),
        "word_count": word_count,
        "table_count": len(doc.tables),
        "styles_used": sorted(styles_used),
    }, indent=2)


def duplicate_document(filename: str, new_filename: str) -> str:
    """Create a copy of a document with a new filename."""
    doc = _load(filename)
    _save(doc, new_filename)
    return json.dumps({"message": f"Document copied to '{new_filename}'."})


def read_document(filename: str, start: int = 0, limit: int = 30) -> str:
    """Read document content with pagination. Each paragraph prefixed with [P<index>|<style>].
    Args:
        filename: Document to read.
        start: 0-based paragraph index to start from (default 0).
        limit: Max paragraphs to return (default 30).
    Returns: Paragraphs as text, plus total count."""
    doc = _load(filename)
    paras = doc.paragraphs
    total = len(paras)
    end = min(start + limit, total)
    lines: list[str] = []
    for i in range(start, end):
        p = paras[i]
        style = p.style.name if p.style else "Normal"
        lines.append(f"[P{i}|{style}] {p.text}")
    header = f"--- Document: {filename} | Paragraphs {start}-{end-1} of {total} ---"
    footer = ""
    if end < total:
        footer = f"\n--- {total - end} more paragraphs. Use read_document('{filename}', start={end}) to continue. ---"
    return header + "\n" + "\n".join(lines) + footer


def read_paragraph(filename: str, paragraph_index: int) -> str:
    """Read a single paragraph by 0-based index. Returns text and style."""
    doc = _load(filename)
    paras = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paras):
        return json.dumps({"error": f"Index {paragraph_index} out of range (0-{len(paras)-1})."})
    p = paras[paragraph_index]
    return json.dumps({
        "index": paragraph_index,
        "style": p.style.name if p.style else "Normal",
        "text": p.text,
        "alignment": str(p.alignment) if p.alignment else "left",
    })


def read_table(filename: str, table_index: int = 0) -> str:
    """Read a table by 0-based index. Returns a JSON 2D array of cell texts."""
    doc = _load(filename)
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        return json.dumps({"error": f"Table index {table_index} out of range (0-{len(tables)-1})."})
    table = tables[table_index]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    return json.dumps({"table_index": table_index, "rows": len(table.rows),
                        "cols": len(table.columns), "data": data}, indent=2)


def search_text(filename: str, query: str, case_sensitive: bool = False) -> str:
    """Search for text in a document. Returns matching paragraphs with indices."""
    doc = _load(filename)
    results = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        q = query if case_sensitive else query.lower()
        t = text if case_sensitive else text.lower()
        if q in t:
            results.append({"paragraph_index": i, "style": p.style.name, "text": text})
    return json.dumps({"query": query, "matches": len(results), "results": results}, indent=2)


def search_and_replace(filename: str, find_text: str, replace_with: str,
                            case_sensitive: bool = False) -> str:
    """Find and replace all occurrences of text in the document."""
    doc = _load(filename)
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if case_sensitive:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_with)
                    count += 1
            else:
                if find_text.lower() in run.text.lower():
                    pattern = re.compile(re.escape(find_text), re.IGNORECASE)
                    count += len(pattern.findall(run.text))
                    run.text = pattern.sub(replace_with, run.text)
    _save(doc, filename)
    return json.dumps({"find": find_text, "replace_with": replace_with, "replacements": count})


def add_paragraph(filename: str, text: str, style: str = "Normal") -> str:
    """Append a new paragraph at the end. Styles: Normal, Heading 1-4, List Bullet, etc."""
    doc = _load(filename)
    doc.add_paragraph(text, style=style)
    _save(doc, filename)
    idx = len(doc.paragraphs) - 1
    return json.dumps({"message": f"Paragraph added at index {idx}.", "index": idx, "style": style})


def insert_paragraph_after(filename: str, text: str, after_index: int,
                                style: str = "Normal") -> str:
    """Insert a new paragraph after the paragraph at the given 0-based index."""
    doc = _load(filename)
    paras = doc.paragraphs
    if after_index < 0 or after_index >= len(paras):
        return json.dumps({"error": f"Index {after_index} out of range."})
    ref_para = paras[after_index]
    new_p = OxmlElement("w:p")
    ref_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, doc)
    new_para.style = style
    new_para.add_run(text)
    _save(doc, filename)
    return json.dumps({"message": f"Paragraph inserted after index {after_index}.",
                        "new_index": after_index + 1})


def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete the paragraph at the given 0-based index."""
    doc = _load(filename)
    paras = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paras):
        return json.dumps({"error": f"Index {paragraph_index} out of range."})
    p_elem = paras[paragraph_index]._p
    p_elem.getparent().remove(p_elem)
    _save(doc, filename)
    return json.dumps({"message": f"Paragraph {paragraph_index} deleted."})


def edit_paragraph_text(filename: str, paragraph_index: int, new_text: str) -> str:
    """Replace the entire text of a paragraph while keeping its style."""
    doc = _load(filename)
    paras = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paras):
        return json.dumps({"error": f"Index {paragraph_index} out of range."})
    para = paras[paragraph_index]
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)
    _save(doc, filename)
    return json.dumps({"message": f"Paragraph {paragraph_index} text updated."})


def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading (level 1-4) at the end of the document."""
    if level < 1 or level > 4:
        return json.dumps({"error": "Heading level must be 1-4."})
    doc = _load(filename)
    doc.add_heading(text, level=level)
    _save(doc, filename)
    idx = len(doc.paragraphs) - 1
    return json.dumps({"message": f"Heading {level} added at index {idx}.", "index": idx})


def add_page_break(filename: str) -> str:
    """Add a page break at the end of the document."""
    doc = _load(filename)
    doc.add_page_break()
    _save(doc, filename)
    return json.dumps({"message": "Page break added."})


def add_table(filename: str, rows: int, cols: int,
                   data: Optional[str] = None) -> str:
    """Add a table at the end. Optional data is a JSON 2D array e.g. '[["A","B"],["1","2"]]'."""
    doc = _load(filename)
    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")
    if data:
        try:
            cell_data = json.loads(data)
            for r, row_data in enumerate(cell_data):
                for c, val in enumerate(row_data):
                    if r < rows and c < cols:
                        table.cell(r, c).text = str(val)
        except (json.JSONDecodeError, IndexError):
            pass
    _save(doc, filename)
    return json.dumps({"message": f"Table ({rows}x{cols}) added.",
                        "table_index": len(doc.tables) - 1})


def add_bullet_list(filename: str, items: str) -> str:
    """Add a bullet list. items is a JSON array e.g. '["First", "Second"]'."""
    doc = _load(filename)
    try:
        item_list = json.loads(items)
    except json.JSONDecodeError:
        return json.dumps({"error": "items must be a JSON array of strings."})
    for item in item_list:
        doc.add_paragraph(str(item), style="List Bullet")
    _save(doc, filename)
    return json.dumps({"message": f"Bullet list with {len(item_list)} items added."})


def add_numbered_list(filename: str, items: str) -> str:
    """Add a numbered list. items is a JSON array e.g. '["First", "Second"]'."""
    doc = _load(filename)
    try:
        item_list = json.loads(items)
    except json.JSONDecodeError:
        return json.dumps({"error": "items must be a JSON array of strings."})
    for item in item_list:
        doc.add_paragraph(str(item), style="List Number")
    _save(doc, filename)
    return json.dumps({"message": f"Numbered list with {len(item_list)} items added."})


def format_text(
    filename: str,
    search_text: str,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    strikethrough: Optional[bool] = None,
    font_size: Optional[float] = None,
    font_name: Optional[str] = None,
    font_color: Optional[str] = None,
    highlight_color: Optional[str] = None,
) -> str:
    """Apply formatting to ALL occurrences of search_text. Colors are hex without '#'."""
    doc = _load(filename)
    fmt = {k: v for k, v in dict(
        bold=bold, italic=italic, underline=underline, strikethrough=strikethrough,
        font_size=font_size, font_name=font_name, font_color=font_color,
        highlight_color=highlight_color,
    ).items() if v is not None}
    formatted_count = 0
    for para in doc.paragraphs:
        if search_text in para.text:
            if _format_matching_text(para, search_text, **fmt):
                formatted_count += 1
    _save(doc, filename)
    return json.dumps({"search_text": search_text, "paragraphs_formatted": formatted_count,
                        "formatting_applied": list(fmt.keys())})


def format_paragraph(
    filename: str,
    paragraph_index: int,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    strikethrough: Optional[bool] = None,
    font_size: Optional[float] = None,
    font_name: Optional[str] = None,
    font_color: Optional[str] = None,
    alignment: Optional[str] = None,
    style: Optional[str] = None,
) -> str:
    """Apply formatting to an entire paragraph. alignment: left/center/right/justify."""
    doc = _load(filename)
    paras = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paras):
        return json.dumps({"error": f"Index {paragraph_index} out of range."})
    para = paras[paragraph_index]
    if alignment and alignment.lower() in ALIGNMENT_MAP:
        para.alignment = ALIGNMENT_MAP[alignment.lower()]
    if style:
        try:
            para.style = style
        except KeyError:
            return json.dumps({"error": f"Unknown style '{style}'."})
    fmt = {k: v for k, v in dict(
        bold=bold, italic=italic, underline=underline, strikethrough=strikethrough,
        font_size=font_size, font_name=font_name, font_color=font_color,
    ).items() if v is not None}
    if fmt:
        for run in para.runs:
            _apply_run_format(run, **fmt)
        if not para.runs and para.text:
            run = para.add_run(para.text)
            _apply_run_format(run, **fmt)
    _save(doc, filename)
    applied = list(fmt.keys())
    if alignment:
        applied.append("alignment")
    if style:
        applied.append("style")
    return json.dumps({"message": f"Paragraph {paragraph_index} formatted.",
                        "formatting_applied": applied})


def set_paragraph_style(filename: str, paragraph_index: int, style: str) -> str:
    """Change paragraph style. Common: Normal, Heading 1-4, List Bullet, Quote, Title."""
    doc = _load(filename)
    paras = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paras):
        return json.dumps({"error": f"Index {paragraph_index} out of range."})
    try:
        paras[paragraph_index].style = style
    except KeyError:
        return json.dumps({"error": f"Unknown style '{style}'."})
    _save(doc, filename)
    return json.dumps({"message": f"Paragraph {paragraph_index} style set to '{style}'."})


def set_paragraph_alignment(filename: str, paragraph_index: int, alignment: str) -> str:
    """Set paragraph alignment: left, center, right, or justify."""
    doc = _load(filename)
    paras = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paras):
        return json.dumps({"error": f"Index {paragraph_index} out of range."})
    a = ALIGNMENT_MAP.get(alignment.lower())
    if a is None:
        return json.dumps({"error": f"Unknown alignment '{alignment}'. Use left/center/right/justify."})
    paras[paragraph_index].alignment = a
    _save(doc, filename)
    return json.dumps({"message": f"Paragraph {paragraph_index} alignment set to '{alignment}'."})


def export_to_text(filename: str) -> str:
    """Export the document as plain text (no formatting)."""
    doc = _load(filename)
    return "\n".join(p.text for p in doc.paragraphs)


ALL_TOOLS = [
    list_documents,
    create_document,
    delete_document,
    get_document_info,
    duplicate_document,
    read_document,
    read_paragraph,
    read_table,
    search_text,
    search_and_replace,
    add_paragraph,
    insert_paragraph_after,
    delete_paragraph,
    edit_paragraph_text,
    add_heading,
    add_page_break,
    add_table,
    add_bullet_list,
    add_numbered_list,
    format_text,
    format_paragraph,
    set_paragraph_style,
    set_paragraph_alignment,
    export_to_text,
]


# ═════════════════════════════════════════════════════════════════════════
#  AGNO AGENT  (Groq LLM)
# ═════════════════════════════════════════════════════════════════════════

AGENT_AVAILABLE = False
root_agent = None

try:
    from agno.agent import Agent
    from agno.models.groq import Groq
    from agno.db.sqlite import SqliteDb

    _DB_FILE = str(Path(__file__).resolve().parent / "agent_sessions.db")
    agent_db = SqliteDb(db_file=_DB_FILE)

    AGENT_MODEL = Groq(id="llama-3.3-70b-versatile")

    AGENT_INSTRUCTION = """\
You are **DocBot**, an intelligent Document Editor Agent for .docx files.

## RULES
1. **Read-only** tools — call immediately, no confirmation needed.
2. **Write/edit/delete** — describe what you will do, then execute. Report: Done.
3. `read_document` is **paginated**: returns 30 paragraphs at a time. Use `start` param to page.
4. Paragraph indices are **0-based** (P0, P1 …).
5. Color hex **without** '#' (e.g. "FF0000"). Table data / list items as JSON strings.
6. If user mentions `[Working on file: X]`, always use file X.
7. Be concise but helpful. Answer questions, summarise, suggest improvements.

## IMPORTANT — TOOL CALLING FORMAT
- NEVER generate function calls as text (e.g. <function=...>).
- ALWAYS use the proper tool-calling API provided by the system.
- If a tool call fails, DO NOT retry it as text — just tell the user what happened.
- You have access to tools like `read_document`, `add_paragraph`, `edit_paragraph_text`, etc.
  Call them through the API, never as inline text.
"""

    root_agent = Agent(
        name="DocBot",
        model=AGENT_MODEL,
        db=agent_db,
        add_history_to_context=True,
        num_history_runs=3,
        store_tool_messages=False,
        description=(
            "DocBot — an intelligent document assistant that can read, discuss, "
            "summarise, question-answer, edit, and format .docx files via "
            "natural-language prompts."
        ),
        instructions=AGENT_INSTRUCTION,
        tools=ALL_TOOLS,
        markdown=True,
    )

    AGENT_AVAILABLE = True
    log.info("AI agent loaded successfully (Agno + Groq).")

except Exception as exc:
    root_agent = None
    AGENT_AVAILABLE = False
    log.warning(f"AI agent unavailable — chat endpoint disabled. Reason: {exc}")


# ═════════════════════════════════════════════════════════════════════════
#  MCP SERVER  (Custom Document Editor MCP — stdio transport)
# ═════════════════════════════════════════════════════════════════════════

MCP_AVAILABLE = False
mcp = None

try:
    from mcp.server.fastmcp import FastMCP

    mcp = FastMCP(
        "Document Editor",
        instructions=(
            "A local document editing server. Manages .docx files stored in "
            f"'{DOCUMENTS_DIR}'. Supports create, read, edit, format, search, "
            "and export operations."
        ),
    )

    @mcp.tool()
    def mcp_list_documents() -> str:
        """List all .docx documents in the document store."""
        return list_documents()

    @mcp.tool()
    def mcp_create_document(title: str) -> str:
        """Create a new blank .docx document with the given title."""
        return create_document(title)

    @mcp.tool()
    def mcp_delete_document(filename: str) -> str:
        """Permanently delete a document from the store."""
        return delete_document(filename)

    @mcp.tool()
    def mcp_get_document_info(filename: str) -> str:
        """Get metadata about a document."""
        return get_document_info(filename)

    @mcp.tool()
    def mcp_read_document(filename: str) -> str:
        """Read the full content of a document."""
        doc = _load(filename)
        lines: list[str] = []
        for i, para in enumerate(doc.paragraphs):
            style = para.style.name if para.style else "Normal"
            lines.append(f"[P{i}|{style}] {para.text}")
        return "\n".join(lines)

    @mcp.tool()
    def mcp_read_paragraph(filename: str, paragraph_index: int) -> str:
        """Read a single paragraph by its index (0-based)."""
        return read_paragraph(filename, paragraph_index)

    @mcp.tool()
    def mcp_read_table(filename: str, table_index: int = 0) -> str:
        """Read a table by its index (0-based)."""
        return read_table(filename, table_index)

    @mcp.tool()
    def mcp_search_text(filename: str, query: str, case_sensitive: bool = False) -> str:
        """Search for text in a document."""
        return search_text(filename, query, case_sensitive)

    @mcp.tool()
    def mcp_search_and_replace(filename: str, find_text: str, replace_with: str,
                               case_sensitive: bool = False) -> str:
        """Find and replace all occurrences of text in a document."""
        return search_and_replace(filename, find_text, replace_with, case_sensitive)

    @mcp.tool()
    def mcp_add_paragraph(filename: str, text: str, style: str = "Normal") -> str:
        """Append a new paragraph at the end of the document."""
        return add_paragraph(filename, text, style)

    @mcp.tool()
    def mcp_insert_paragraph_after(filename: str, text: str, after_index: int,
                                   style: str = "Normal") -> str:
        """Insert a new paragraph after the paragraph at the given index."""
        return insert_paragraph_after(filename, text, after_index, style)

    @mcp.tool()
    def mcp_delete_paragraph(filename: str, paragraph_index: int) -> str:
        """Delete the paragraph at the given index (0-based)."""
        return delete_paragraph(filename, paragraph_index)

    @mcp.tool()
    def mcp_edit_paragraph_text(filename: str, paragraph_index: int, new_text: str) -> str:
        """Replace the entire text of a paragraph while keeping its style."""
        return edit_paragraph_text(filename, paragraph_index, new_text)

    @mcp.tool()
    def mcp_add_heading(filename: str, text: str, level: int = 1) -> str:
        """Add a heading (level 1-4) at the end of the document."""
        return add_heading(filename, text, level)

    @mcp.tool()
    def mcp_add_page_break(filename: str) -> str:
        """Add a page break at the end of the document."""
        return add_page_break(filename)

    @mcp.tool()
    def mcp_add_table(filename: str, rows: int, cols: int,
                      data: str | None = None) -> str:
        """Add a table at the end of the document. Optional data is a JSON 2D array."""
        return add_table(filename, rows, cols, data)

    @mcp.tool()
    def mcp_add_bullet_list(filename: str, items: str) -> str:
        """Add a bullet list. items is a JSON array of strings."""
        return add_bullet_list(filename, items)

    @mcp.tool()
    def mcp_add_numbered_list(filename: str, items: str) -> str:
        """Add a numbered list. items is a JSON array of strings."""
        return add_numbered_list(filename, items)

    @mcp.tool()
    def mcp_format_text(
        filename: str, search_text: str,
        bold: bool | None = None, italic: bool | None = None,
        underline: bool | None = None, strikethrough: bool | None = None,
        font_size: float | None = None, font_name: str | None = None,
        font_color: str | None = None, highlight_color: str | None = None,
    ) -> str:
        """Apply formatting to ALL occurrences of search_text. Colors hex without '#'."""
        return format_text(filename, search_text, bold, italic, underline,
                                strikethrough, font_size, font_name, font_color, highlight_color)

    @mcp.tool()
    def mcp_format_paragraph(
        filename: str, paragraph_index: int,
        bold: bool | None = None, italic: bool | None = None,
        underline: bool | None = None, strikethrough: bool | None = None,
        font_size: float | None = None, font_name: str | None = None,
        font_color: str | None = None, alignment: str | None = None,
        style: str | None = None,
    ) -> str:
        """Apply formatting to an entire paragraph."""
        return format_paragraph(filename, paragraph_index, bold, italic, underline,
                                     strikethrough, font_size, font_name, font_color, alignment, style)

    @mcp.tool()
    def mcp_set_paragraph_style(filename: str, paragraph_index: int, style: str) -> str:
        """Change the style of a paragraph."""
        return set_paragraph_style(filename, paragraph_index, style)

    @mcp.tool()
    def mcp_set_paragraph_alignment(filename: str, paragraph_index: int, alignment: str) -> str:
        """Set paragraph alignment: left, center, right, or justify."""
        return set_paragraph_alignment(filename, paragraph_index, alignment)

    @mcp.tool()
    def mcp_export_to_text(filename: str) -> str:
        """Export the document as plain text."""
        return export_to_text(filename)

    @mcp.tool()
    def mcp_duplicate_document(filename: str, new_filename: str) -> str:
        """Create a copy of a document with a new filename."""
        return duplicate_document(filename, new_filename)

    MCP_AVAILABLE = True
    log.info("MCP server loaded successfully.")

except Exception as exc:
    MCP_AVAILABLE = False
    log.warning(f"MCP server unavailable. Reason: {exc}")


# ═════════════════════════════════════════════════════════════════════════
#  FASTAPI APP
# ═════════════════════════════════════════════════════════════════════════

app = FastAPI(
    title="Document Editor Agent API",
    description="Unified REST API for the Document Editor Agent with MCP Server.",
    version="3.0.0",
    docs_url="/docs",
    redoc_url="/redoc",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Pydantic Models ─────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    user_id: str = "default_user"
    session_id: str | None = None

class ChatResponse(BaseModel):
    response: str
    user_id: str
    session_id: str

class CreateDocumentRequest(BaseModel):
    title: str

class DuplicateRequest(BaseModel):
    new_filename: str

class AddParagraphRequest(BaseModel):
    text: str
    style: str = "Normal"

class EditParagraphRequest(BaseModel):
    new_text: str

class InsertParagraphRequest(BaseModel):
    text: str
    style: str = "Normal"

class AddHeadingRequest(BaseModel):
    text: str
    level: int = 1

class AddTableRequest(BaseModel):
    rows: int
    cols: int
    data: list[list[str]] | None = None

class ListItemsRequest(BaseModel):
    items: list[str]

class SearchRequest(BaseModel):
    query: str
    case_sensitive: bool = False

class ReplaceRequest(BaseModel):
    find_text: str
    replace_with: str
    case_sensitive: bool = False

class FormatParagraphRequest(BaseModel):
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    strikethrough: bool | None = None
    font_size: float | None = None
    font_name: str | None = None
    font_color: str | None = None
    alignment: str | None = None
    style: str | None = None

class FormatTextRequest(BaseModel):
    search_text: str
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    strikethrough: bool | None = None
    font_size: float | None = None
    font_name: str | None = None
    font_color: str | None = None
    highlight_color: str | None = None

class SetStyleRequest(BaseModel):
    style: str

class SetAlignmentRequest(BaseModel):
    alignment: str


# ═════════════════════════════════════════════════════════════════════════
#  HEALTH
# ═════════════════════════════════════════════════════════════════════════

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "agent_available": AGENT_AVAILABLE,
        "mcp_available": MCP_AVAILABLE,
        "agent_name": root_agent.name if AGENT_AVAILABLE else None,
        "documents_dir": str(DOCUMENTS_DIR),
        "document_count": len(list(DOCUMENTS_DIR.glob("*.docx"))),
    }


# ═════════════════════════════════════════════════════════════════════════
#  AI CHAT
# ═════════════════════════════════════════════════════════════════════════

MAX_RETRIES = 3
INITIAL_BACKOFF = 1


@app.post("/chat", response_model=ChatResponse)
async def chat(req: ChatRequest):
    """Send a natural-language message to the AI document-editor agent."""
    log.info(f"Incoming Chat Request [User: {req.user_id}, Session: {req.session_id}]: {req.message[:100]}...")
    
    if not AGENT_AVAILABLE:
        log.error("Chat rejected: AI agent not available.")
        raise HTTPException(status_code=503, detail="AI agent is not available. Check GROQ_API_KEY.")

    user_id = req.user_id
    session_id = req.session_id or str(uuid.uuid4())
    last_err: str = ""

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            log.info(f"Generating agent response (Attempt {attempt}/{MAX_RETRIES})...")
            start_time = time.time()
            run_output = await root_agent.arun(
                input=req.message,
                user_id=user_id,
                session_id=session_id,
            )
            elapsed = time.time() - start_time
            log.info(f"Agent response generated in {elapsed:.2f}s.")
            response_text = run_output.content if run_output.content else "(no response)"

            # Check if the model returned a tool_use_failed error as text
            is_tool_fail_in_text = (
                "tool_use_failed" in response_text
                or "failed_generation" in response_text
                or "Failed to call a function" in response_text
            )
            if is_tool_fail_in_text and attempt < MAX_RETRIES:
                session_id = str(uuid.uuid4())
                log.warning(
                    f"tool_use_failed in response text (attempt {attempt}/{MAX_RETRIES}), "
                    f"retrying with fresh session {session_id[:8]}…"
                )
                await asyncio.sleep(0.5)
                continue

            if is_tool_fail_in_text:
                # All retries exhausted, return a friendly message
                return ChatResponse(
                    response="I had trouble executing the tool. Please rephrase your request or try again.",
                    user_id=user_id,
                    session_id=session_id,
                )

            if "rate_limit_exceeded" in response_text or "rate limit reached" in response_text.lower():
                raise HTTPException(status_code=429, detail="Groq API rate limit reached.")

            return ChatResponse(response=response_text, user_id=user_id, session_id=session_id)

        except HTTPException:
            raise
        except Exception as e:
            err_msg = str(e)
            last_err = err_msg
            is_rate_limit = "429" in err_msg or "rate" in err_msg.lower()
            is_tool_fail = "tool_use_failed" in err_msg or "Failed to call a function" in err_msg

            if is_tool_fail and attempt < MAX_RETRIES:
                session_id = str(uuid.uuid4())
                log.warning(
                    f"tool_use_failed exception (attempt {attempt}/{MAX_RETRIES}), "
                    f"retrying with fresh session {session_id[:8]}…"
                )
                await asyncio.sleep(0.5)
                continue

            if is_rate_limit and attempt < MAX_RETRIES:
                wait = INITIAL_BACKOFF * (2 ** (attempt - 1))
                log.warning(f"Rate-limit (attempt {attempt}/{MAX_RETRIES}), retrying in {wait}s…")
                await asyncio.sleep(wait)
                continue

            if is_rate_limit:
                raise HTTPException(status_code=429, detail="Groq API rate limit exceeded after retries.")
            if is_tool_fail:
                return ChatResponse(
                    response="I had trouble executing the tool. Please rephrase your request or try again.",
                    user_id=user_id,
                    session_id=session_id,
                )

            raise HTTPException(status_code=500, detail=f"Agent error: {err_msg}")

    return ChatResponse(
        response="I encountered repeated errors. Please try a New Chat Session and rephrase your request.",
        user_id=user_id,
        session_id=session_id,
    )


@app.post("/sessions")
async def create_session():
    return {"session_id": str(uuid.uuid4())}


# ═════════════════════════════════════════════════════════════════════════
#  DOCUMENT MANAGEMENT ENDPOINTS
# ═════════════════════════════════════════════════════════════════════════

@app.get("/documents")
async def api_list_documents():
    docs = []
    for f in sorted(DOCUMENTS_DIR.glob("*.docx")):
        stat = f.stat()
        entry = {"filename": f.name, "size_bytes": stat.st_size, "modified": stat.st_mtime}
        # Check if file is a valid docx (zip archive)
        try:
            with zipfile.ZipFile(str(f), "r"):
                pass
            entry["corrupted"] = False
        except (zipfile.BadZipFile, Exception):
            entry["corrupted"] = True
        docs.append(entry)
    return docs


@app.post("/documents")
async def api_create_document(req: CreateDocumentRequest):
    safe_title = re.sub(r'[<>:"/\\|?*]', "_", req.title)
    filename = f"{safe_title}.docx"
    doc = Document()
    doc.add_heading(req.title, level=1)
    _save(doc, filename)
    return {"filename": filename, "message": f"Document '{req.title}' created."}


@app.post("/documents/upload")
async def api_upload_document(file: UploadFile = File(...)):
    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are accepted.")
    dest = DOCUMENTS_DIR / Path(file.filename).name
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)
    # Validate the uploaded file is a real docx
    try:
        Document(str(dest))
    except (zipfile.BadZipFile, Exception) as exc:
        dest.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Uploaded file is not a valid .docx document: {exc}")
    stat = dest.stat()
    return {"message": f"Uploaded '{dest.name}'.", "filename": dest.name, "size_bytes": stat.st_size}


@app.get("/documents/{filename}")
async def api_get_document_info(filename: str):
    try:
        doc = _load(filename)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")
    except zipfile.BadZipFile:
        raise HTTPException(status_code=422, detail=f"Document '{filename}' is corrupted and cannot be opened.")
    paras = doc.paragraphs
    word_count = sum(len(p.text.split()) for p in paras)
    styles_used = sorted({p.style.name for p in paras if p.style})
    return {
        "filename": filename,
        "paragraph_count": len(paras),
        "word_count": word_count,
        "table_count": len(doc.tables),
        "styles_used": styles_used,
    }


@app.get("/documents/{filename}/download")
async def api_download_document(filename: str):
    path = _resolve(filename)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")
    return FileResponse(
        path=str(path), filename=path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/documents/{filename}/content")
async def api_get_document_content(filename: str, start: int = 0, limit: int = 100):
    try:
        doc = _load(filename)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")
    except zipfile.BadZipFile:
        raise HTTPException(status_code=422, detail=f"Document '{filename}' is corrupted and cannot be read.")
    paras = doc.paragraphs
    total = len(paras)
    end = min(start + limit, total)
    paragraphs = []
    for i in range(start, end):
        p = paras[i]
        paragraphs.append({"index": i, "text": p.text, "style": p.style.name if p.style else "Normal"})
    return {
        "filename": filename, "total_paragraphs": total,
        "start": start, "end": end - 1, "has_more": end < total,
        "paragraphs": paragraphs,
    }


@app.get("/documents/{filename}/preview", response_class=HTMLResponse)
async def api_get_document_preview(filename: str):
    try:
        doc = _load(filename)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")
    except zipfile.BadZipFile:
        raise HTTPException(status_code=422, detail=f"Document '{filename}' is corrupted and cannot be previewed.")
    return _docx_to_html(doc)


@app.get("/documents/{filename}/snapshot")
async def api_get_document_snapshot(filename: str):
    """Quick snapshot for before/after comparison — paragraph count, word count, last 5 paragraphs."""
    try:
        doc = _load(filename)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")
    except zipfile.BadZipFile:
        raise HTTPException(status_code=422, detail=f"Document '{filename}' is corrupted.")
    paras = doc.paragraphs
    total = len(paras)
    word_count = sum(len(p.text.split()) for p in paras)
    last_5 = []
    start = max(0, total - 5)
    for i in range(start, total):
        p = paras[i]
        last_5.append({
            "index": i,
            "style": p.style.name if p.style else "Normal",
            "text": p.text[:200],
        })
    return {
        "filename": filename,
        "paragraph_count": total,
        "word_count": word_count,
        "table_count": len(doc.tables),
        "last_paragraphs": last_5,
    }


@app.delete("/documents/{filename}")
async def api_delete_document(filename: str):
    path = _resolve(filename)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")
    path.unlink()
    return {"message": f"Document '{filename}' deleted."}


@app.post("/documents/{filename}/duplicate")
async def api_duplicate_document(filename: str, req: DuplicateRequest):
    try:
        doc = _load(filename)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")
    _save(doc, req.new_filename)
    return {"message": f"Document copied to '{req.new_filename}'."}


# ═════════════════════════════════════════════════════════════════════════
#  DIRECT EDITING ENDPOINTS
# ═════════════════════════════════════════════════════════════════════════

def _ensure_doc(filename: str) -> Document:
    try:
        return _load(filename)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")


@app.post("/documents/{filename}/paragraphs")
async def api_add_paragraph(filename: str, req: AddParagraphRequest):
    doc = _ensure_doc(filename)
    doc.add_paragraph(req.text, style=req.style)
    _save(doc, filename)
    idx = len(doc.paragraphs) - 1
    return {"message": f"Paragraph added at index {idx}.", "index": idx, "style": req.style}


@app.put("/documents/{filename}/paragraphs/{index}")
async def api_edit_paragraph(filename: str, index: int, req: EditParagraphRequest):
    doc = _ensure_doc(filename)
    paras = doc.paragraphs
    if index < 0 or index >= len(paras):
        raise HTTPException(status_code=400, detail=f"Index {index} out of range (0-{len(paras)-1}).")
    para = paras[index]
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = req.new_text
    else:
        para.add_run(req.new_text)
    _save(doc, filename)
    return {"message": f"Paragraph {index} updated."}


@app.delete("/documents/{filename}/paragraphs/{index}")
async def api_delete_paragraph(filename: str, index: int):
    doc = _ensure_doc(filename)
    paras = doc.paragraphs
    if index < 0 or index >= len(paras):
        raise HTTPException(status_code=400, detail=f"Index {index} out of range.")
    p_elem = paras[index]._p
    p_elem.getparent().remove(p_elem)
    _save(doc, filename)
    return {"message": f"Paragraph {index} deleted."}


@app.post("/documents/{filename}/paragraphs/{index}/insert")
async def api_insert_paragraph_after(filename: str, index: int, req: InsertParagraphRequest):
    doc = _ensure_doc(filename)
    paras = doc.paragraphs
    if index < 0 or index >= len(paras):
        raise HTTPException(status_code=400, detail=f"Index {index} out of range.")
    ref_para = paras[index]
    new_p = OxmlElement("w:p")
    ref_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, doc)
    new_para.style = req.style
    new_para.add_run(req.text)
    _save(doc, filename)
    return {"message": f"Paragraph inserted after index {index}.", "new_index": index + 1}


@app.post("/documents/{filename}/headings")
async def api_add_heading(filename: str, req: AddHeadingRequest):
    if req.level < 1 or req.level > 4:
        raise HTTPException(status_code=400, detail="Heading level must be 1-4.")
    doc = _ensure_doc(filename)
    doc.add_heading(req.text, level=req.level)
    _save(doc, filename)
    idx = len(doc.paragraphs) - 1
    return {"message": f"Heading {req.level} added at index {idx}.", "index": idx}


@app.post("/documents/{filename}/tables")
async def api_add_table(filename: str, req: AddTableRequest):
    doc = _ensure_doc(filename)
    table = doc.add_table(rows=req.rows, cols=req.cols, style="Table Grid")
    if req.data:
        for r, row_data in enumerate(req.data):
            for c, val in enumerate(row_data):
                if r < req.rows and c < req.cols:
                    table.cell(r, c).text = str(val)
    _save(doc, filename)
    return {"message": f"Table ({req.rows}x{req.cols}) added.", "table_index": len(doc.tables) - 1}


@app.post("/documents/{filename}/lists/bullet")
async def api_add_bullet_list(filename: str, req: ListItemsRequest):
    doc = _ensure_doc(filename)
    for item in req.items:
        doc.add_paragraph(str(item), style="List Bullet")
    _save(doc, filename)
    return {"message": f"Bullet list with {len(req.items)} items added."}


@app.post("/documents/{filename}/lists/numbered")
async def api_add_numbered_list(filename: str, req: ListItemsRequest):
    doc = _ensure_doc(filename)
    for item in req.items:
        doc.add_paragraph(str(item), style="List Number")
    _save(doc, filename)
    return {"message": f"Numbered list with {len(req.items)} items added."}


@app.post("/documents/{filename}/search")
async def api_search_text(filename: str, req: SearchRequest):
    doc = _ensure_doc(filename)
    results = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        q = req.query if req.case_sensitive else req.query.lower()
        t = text if req.case_sensitive else text.lower()
        if q in t:
            results.append({"paragraph_index": i, "style": p.style.name, "text": text})
    return {"query": req.query, "matches": len(results), "results": results}


@app.post("/documents/{filename}/replace")
async def api_search_and_replace(filename: str, req: ReplaceRequest):
    doc = _ensure_doc(filename)
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if req.case_sensitive:
                if req.find_text in run.text:
                    run.text = run.text.replace(req.find_text, req.replace_with)
                    count += 1
            else:
                if req.find_text.lower() in run.text.lower():
                    pattern = re.compile(re.escape(req.find_text), re.IGNORECASE)
                    count += len(pattern.findall(run.text))
                    run.text = pattern.sub(req.replace_with, run.text)
    _save(doc, filename)
    return {"find": req.find_text, "replace_with": req.replace_with, "replacements": count}


@app.put("/documents/{filename}/paragraphs/{index}/format")
async def api_format_paragraph(filename: str, index: int, req: FormatParagraphRequest):
    doc = _ensure_doc(filename)
    paras = doc.paragraphs
    if index < 0 or index >= len(paras):
        raise HTTPException(status_code=400, detail=f"Index {index} out of range.")
    para = paras[index]
    if req.alignment and req.alignment.lower() in ALIGNMENT_MAP:
        para.alignment = ALIGNMENT_MAP[req.alignment.lower()]
    if req.style:
        try:
            para.style = req.style
        except KeyError:
            raise HTTPException(status_code=400, detail=f"Unknown style '{req.style}'.")
    fmt = {k: v for k, v in dict(
        bold=req.bold, italic=req.italic, underline=req.underline,
        strikethrough=req.strikethrough, font_size=req.font_size,
        font_name=req.font_name, font_color=req.font_color,
    ).items() if v is not None}
    if fmt:
        for run in para.runs:
            _apply_run_format(run, **fmt)
        if not para.runs and para.text:
            run = para.add_run(para.text)
            _apply_run_format(run, **fmt)
    _save(doc, filename)
    applied = list(fmt.keys())
    if req.alignment:
        applied.append("alignment")
    if req.style:
        applied.append("style")
    return {"message": f"Paragraph {index} formatted.", "formatting_applied": applied}


@app.put("/documents/{filename}/paragraphs/{index}/style")
async def api_set_paragraph_style(filename: str, index: int, req: SetStyleRequest):
    doc = _ensure_doc(filename)
    paras = doc.paragraphs
    if index < 0 or index >= len(paras):
        raise HTTPException(status_code=400, detail=f"Index {index} out of range.")
    try:
        paras[index].style = req.style
    except KeyError:
        raise HTTPException(status_code=400, detail=f"Unknown style '{req.style}'.")
    _save(doc, filename)
    return {"message": f"Paragraph {index} style set to '{req.style}'."}


@app.put("/documents/{filename}/paragraphs/{index}/align")
async def api_set_paragraph_alignment(filename: str, index: int, req: SetAlignmentRequest):
    doc = _ensure_doc(filename)
    paras = doc.paragraphs
    if index < 0 or index >= len(paras):
        raise HTTPException(status_code=400, detail=f"Index {index} out of range.")
    a = ALIGNMENT_MAP.get(req.alignment.lower())
    if a is None:
        raise HTTPException(status_code=400, detail="Use: left/center/right/justify.")
    paras[index].alignment = a
    _save(doc, filename)
    return {"message": f"Paragraph {index} alignment set to '{req.alignment}'."}


@app.post("/documents/{filename}/format-text")
async def api_format_text(filename: str, req: FormatTextRequest):
    doc = _ensure_doc(filename)
    fmt = {k: v for k, v in dict(
        bold=req.bold, italic=req.italic, underline=req.underline,
        strikethrough=req.strikethrough, font_size=req.font_size,
        font_name=req.font_name, font_color=req.font_color,
        highlight_color=req.highlight_color,
    ).items() if v is not None}
    count = 0
    for para in doc.paragraphs:
        if req.search_text in para.text:
            if _format_matching_text(para, req.search_text, **fmt):
                count += 1
    _save(doc, filename)
    return {"search_text": req.search_text, "paragraphs_formatted": count, "formatting_applied": list(fmt.keys())}


# ═════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import sys
    import uvicorn

    # If --mcp flag, run MCP server on stdio
    if "--mcp" in sys.argv:
        if MCP_AVAILABLE and mcp:
            print("Starting MCP server on stdio...")
            mcp.run(transport="stdio")
        else:
            print("MCP server is not available.")
            sys.exit(1)
    else:
        print("=" * 60)
        print("  Document Editor Agent API")
        print(f"  Agent:  {'Loaded' if AGENT_AVAILABLE else 'Unavailable'}")
        print(f"  MCP:    {'Loaded' if MCP_AVAILABLE else 'Unavailable'}")
        print(f"  Docs:   {DOCUMENTS_DIR}")
        print(f"  URL:    http://localhost:8080")
        print(f"  Swagger: http://localhost:8080/docs")
        print("=" * 60)
        uvicorn.run("api:app", host="0.0.0.0", port=8080, reload=True)
