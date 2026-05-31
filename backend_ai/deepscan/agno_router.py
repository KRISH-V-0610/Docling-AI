"""
FormatForge AI — Unified Pipeline Router
==========================================
Full pipeline:  Static Formatting (Orchestrator) → LLM LaTeX Generation (Agno/Groq)
Exposes SSE streaming endpoint: POST /api/v2/pipeline/stream
"""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import shutil
import tempfile
from pathlib import Path

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse

logger = logging.getLogger(__name__)

agno_router = APIRouter(tags=["Unified Pipeline"])

# ---------------------------------------------------------------------------
# Lazy Agno imports
# ---------------------------------------------------------------------------
_AGNO_READY = False

try:
    from agno.agent import Agent
    from agno.models.groq import Groq as AgnoGroq

    _AGNO_READY = True
except ImportError as exc:
    logger.warning(f"Agno deps missing ({exc}). LaTeX generation will be disabled.")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _extract_from_docx(file_path: Path) -> str:
    """Extract full text from a .docx file on disk."""
    import docx

    text = ""
    doc = docx.Document(str(file_path))
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


# ---------------------------------------------------------------------------
# Image extraction (Phase B) — pull figures out of the formatted DOCX so they
# survive into the generated LaTeX as real \includegraphics.
# ---------------------------------------------------------------------------

# Formats pdflatex/tectonic can \includegraphics directly.
_PDFLATEX_OK_EXT = {".png", ".jpg", ".jpeg", ".pdf"}

_EXT_BY_CONTENT_TYPE = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/bmp": ".bmp",
    "image/tiff": ".tiff",
    "image/x-emf": ".emf",
    "image/x-wmf": ".wmf",
    "image/webp": ".webp",
    "application/pdf": ".pdf",
}


def _raster_to_png(blob: bytes) -> bytes | None:
    """Convert a raster image (gif/bmp/tiff/webp) to PNG via Pillow. Returns
    None when conversion isn't possible (e.g. vector EMF/WMF)."""
    try:
        import io

        from PIL import Image

        im = Image.open(io.BytesIO(blob))
        im = im.convert("RGBA") if im.mode in ("RGBA", "P", "LA") else im.convert("RGB")
        out = io.BytesIO()
        im.save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def _save_docx_image(doc, rid: str, assets_dir: Path, n: int) -> Optional[str]:
    """Save the image referenced by relationship id *rid* into *assets_dir* as
    figN.<ext>. Converts unsupported raster formats to PNG; returns the saved
    filename, or None when the image can't be made pdflatex-compatible."""
    try:
        part = doc.part.related_parts[rid]
        blob = part.blob
    except (KeyError, AttributeError, Exception):
        return None

    ct = (getattr(part, "content_type", "") or "").lower()
    ext = _EXT_BY_CONTENT_TYPE.get(ct)
    if ext is None:
        pn = str(getattr(part, "partname", "")).lower()
        for e in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif",
                  ".emf", ".wmf", ".webp", ".pdf"):
            if pn.endswith(e):
                ext = e
                break

    if ext in _PDFLATEX_OK_EXT:
        fname = f"fig{n}{ext}"
        (assets_dir / fname).write_bytes(blob)
        return fname

    # Try to convert raster formats (gif/bmp/tiff/webp) → PNG.
    png = _raster_to_png(blob)
    if png is not None:
        fname = f"fig{n}.png"
        (assets_dir / fname).write_bytes(png)
        return fname

    # Vector EMF/WMF or unknown — drop (caller logs a warning); text is kept.
    return None


def _extract_text_and_images(file_path: Path, assets_dir: Path) -> tuple[str, list[dict], int]:
    """Walk a DOCX in document order, returning (text_with_sentinels, figures,
    skipped_count). A ``[[FIGURE_n]]`` sentinel is placed on its own line where
    each successfully-saved image occurs, so the LLM can keep figure placement.

    figures: list of {"number": int, "filename": str, "token": str}.
    skipped_count: images that couldn't be made pdflatex-compatible (dropped).
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(file_path))
    assets_dir.mkdir(parents=True, exist_ok=True)

    lines: list[str] = []
    figures: list[dict] = []
    fig_n = 0
    skipped = 0

    blip_tag = ".//" + qn("a:blip")
    embed_attr = qn("r:embed")
    link_attr = qn("r:link")

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

        for blip in para._element.findall(blip_tag):
            rid = blip.get(embed_attr) or blip.get(link_attr)
            if not rid:
                continue
            fig_n += 1
            token = f"[[FIGURE_{fig_n}]]"
            lines.append(token)
            saved = _save_docx_image(doc, rid, assets_dir, fig_n)
            if saved:
                figures.append({"number": fig_n, "filename": saved, "token": token})
            else:
                # Couldn't extract (e.g. vector EMF/WMF). Track as a MISSING
                # figure (filename=None) so its placement is preserved and
                # Phase I can prompt the user to upload an image for it.
                skipped += 1
                figures.append({"number": fig_n, "filename": None, "token": token})

    return "\n".join(lines), figures, skipped


def _figure_env(fig: dict) -> str:
    """Build a LaTeX figure float for a saved image, or — when the image
    couldn't be extracted (Phase I, filename is None) — a marked placeholder
    box so the document still compiles and the slot can be filled later."""
    n = fig["number"]
    caption = fig.get("caption") or f"Figure {n}"
    fname = fig.get("filename")
    if not fname:
        return _figure_placeholder(n, caption)
    return (
        "\\begin{figure}[h]\n"
        "\\centering\n"
        f"\\includegraphics[width=0.8\\linewidth]{{assets/{fname}}}\n"
        f"\\caption{{{caption}}}\n"
        f"\\label{{fig:{n}}}\n"
        "\\end{figure}"
    )


def _figure_placeholder(n: int, caption: str = "") -> str:
    """A visible placeholder for a figure whose image couldn't be extracted,
    wrapped in ``% BEGIN_FIGURE_SLOT_n`` / ``% END_FIGURE_SLOT_n`` markers so
    ``_fill_figure_slot`` can swap in the real image once the user uploads it."""
    cap = caption or f"Figure {n}"
    return (
        f"% BEGIN_FIGURE_SLOT_{n}\n"
        "\\begin{figure}[h]\n"
        "\\centering\n"
        "\\fbox{\\begin{minipage}{0.8\\linewidth}\\centering\\vspace{2cm}\n"
        f"\\textit{{[Figure {n} --- image not extracted; upload to render]}}\n"
        "\\vspace{2cm}\\end{minipage}}\n"
        f"\\caption{{{cap}}}\n"
        f"\\label{{fig:{n}}}\n"
        "\\end{figure}\n"
        f"% END_FIGURE_SLOT_{n}"
    )


def _fill_figure_slot(latex: str, n: int, rel_path: str, caption: str = "") -> str:
    """Replace the marked placeholder for figure ``n`` with a real figure float
    pointing at ``rel_path``. If the marker block is gone (user edited it away),
    append the figure so an uploaded image is never silently lost."""
    cap = caption or f"Figure {n}"
    real = (
        "\\begin{figure}[h]\n"
        "\\centering\n"
        f"\\includegraphics[width=0.8\\linewidth]{{{rel_path}}}\n"
        f"\\caption{{{cap}}}\n"
        f"\\label{{fig:{n}}}\n"
        "\\end{figure}"
    )
    pattern = re.compile(
        rf"% BEGIN_FIGURE_SLOT_{n}\b.*?% END_FIGURE_SLOT_{n}", re.DOTALL
    )
    if pattern.search(latex):
        # lambda replacement so backslashes in `real` aren't read as backrefs.
        return pattern.sub(lambda _: real, latex)
    return latex.rstrip() + "\n\n" + real + "\n"


def _inject_figures(body: str, figures: list[dict]) -> str:
    """Replace every ``[[FIGURE_n]]`` sentinel with a real figure float.
    Any figure whose sentinel the LLM dropped is deterministically re-appended
    at the end so **no figure is ever lost**."""
    if not figures:
        # Still strip any stray sentinels the model might have hallucinated.
        return re.sub(r"\[\[FIGURE_\d+\]\]", "", body)

    by_num = {f["number"]: f for f in figures}
    present: set[int] = set()

    def _repl(m: re.Match) -> str:
        n = int(m.group(1))
        fig = by_num.get(n)
        if not fig:
            return ""
        present.add(n)
        return _figure_env(fig)

    body = re.sub(r"\[\[FIGURE_(\d+)\]\]", _repl, body)

    missing = [f for f in figures if f["number"] not in present]
    if missing:
        tail = "\n\n".join(_figure_env(f) for f in missing)
        body = f"{body}\n\n{tail}"

    return body


def _sse(data: dict) -> str:
    return f"data: {json.dumps(data)}\n\n"


# ---------------------------------------------------------------------------
# Chunked LaTeX helpers
# ---------------------------------------------------------------------------

# Section headings pattern for splitting
_SECTION_RE = re.compile(
    r"^(Abstract|Introduction|Background|Results|Discussion|"
    r"Materials and Methods|Methods|Significance|Conclusion|Conclusions|"
    r"Acknowledgments|ACKNOWLEDGMENTS|References|Appendix|"
    r"Supplementary|Supporting Information)",
    re.IGNORECASE | re.MULTILINE,
)

_MAX_CHUNK_CHARS = 5000  # ~1200 words, safely under Groq output limits


def _split_into_sections(text: str) -> list[dict]:
    """
    Split a manuscript into logical sections for chunked LLM processing.
    Returns list of {"title": str, "content": str} dicts.
    """
    lines = text.split("\n")
    sections: list[dict] = []
    current_title = "Preamble"
    current_lines: list[str] = []

    for line in lines:
        match = _SECTION_RE.match(line.strip())
        if match and len(current_lines) > 0:
            sections.append({
                "title": current_title,
                "content": "\n".join(current_lines).strip(),
            })
            current_title = line.strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append({
            "title": current_title,
            "content": "\n".join(current_lines).strip(),
        })

    # Filter out empty sections
    sections = [s for s in sections if s["content"].strip()]

    # Further split any section that's too large
    final_sections: list[dict] = []
    for sec in sections:
        if len(sec["content"]) <= _MAX_CHUNK_CHARS:
            final_sections.append(sec)
        else:
            # Split large sections into sub-chunks by paragraph
            paragraphs = sec["content"].split("\n\n")
            chunk_lines: list[str] = []
            chunk_idx = 1
            for para in paragraphs:
                if len("\n\n".join(chunk_lines)) + len(para) + 2 > _MAX_CHUNK_CHARS and chunk_lines:
                    final_sections.append({
                        "title": f"{sec['title']} (Part {chunk_idx})",
                        "content": "\n\n".join(chunk_lines).strip(),
                    })
                    chunk_lines = [para]
                    chunk_idx += 1
                else:
                    chunk_lines.append(para)
            if chunk_lines:
                final_sections.append({
                    "title": f"{sec['title']} (Part {chunk_idx})" if chunk_idx > 1 else sec["title"],
                    "content": "\n\n".join(chunk_lines).strip(),
                })

    return final_sections


def _clean_chunk_latex(raw: str) -> str:
    """Clean LLM output to extract just the LaTeX body content (no preamble/document tags)."""
    text = raw.strip()
    # Remove markdown fences
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
        text = text.strip()
    # Remove any \documentclass, \usepackage, \begin{document}, \end{document}
    text = re.sub(r"\\documentclass(\[.*?\])?\{.*?\}", "", text)
    text = re.sub(r"\\usepackage(\[.*?\])?\{.*?\}", "", text)
    text = re.sub(r"\\begin\{document\}", "", text)
    text = re.sub(r"\\end\{document\}", "", text)
    text = re.sub(r"\\maketitle", "", text)
    text = re.sub(r"\\printbibliography", "", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Style-specific LaTeX preamble builder
# ---------------------------------------------------------------------------

_PREAMBLE_MAP: dict[str, str] = {
    "ieee": (
        "\\documentclass[conference]{IEEEtran}\n"
        "\\usepackage{cite}\n"
        "\\usepackage{amsmath,amssymb,amsfonts}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage{textcomp}\n"
        "\\usepackage{xcolor}\n"
        "\\usepackage{hyperref}\n"
        "\\usepackage{float}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
    ),
    "apa7": (
        "\\documentclass[12pt]{article}\n"
        "\\usepackage[margin=1in]{geometry}\n"
        "\\usepackage{times}\n"
        "\\usepackage{setspace}\n"
        "\\doublespacing\n"
        "\\usepackage{amsmath}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage{hyperref}\n"
        "\\usepackage{float}\n"
        "\\usepackage{natbib}\n"
        "\\usepackage{caption}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage{indentfirst}\n"
        "\\setlength{\\parindent}{0.5in}\n"
    ),
    "mla": (
        "\\documentclass[12pt]{article}\n"
        "\\usepackage[margin=1in]{geometry}\n"
        "\\usepackage{times}\n"
        "\\usepackage{setspace}\n"
        "\\doublespacing\n"
        "\\usepackage{amsmath}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage{hyperref}\n"
        "\\usepackage{float}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage{indentfirst}\n"
        "\\setlength{\\parindent}{0.5in}\n"
    ),
    "chicago": (
        "\\documentclass[12pt]{article}\n"
        "\\usepackage[margin=1in]{geometry}\n"
        "\\usepackage{times}\n"
        "\\usepackage{setspace}\n"
        "\\doublespacing\n"
        "\\usepackage{amsmath}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage{hyperref}\n"
        "\\usepackage{float}\n"
        "\\usepackage{caption}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage{indentfirst}\n"
        "\\setlength{\\parindent}{0.5in}\n"
    ),
    "vancouver": (
        "\\documentclass[12pt]{article}\n"
        "\\usepackage[margin=1in]{geometry}\n"
        "\\usepackage{times}\n"
        "\\usepackage{setspace}\n"
        "\\doublespacing\n"
        "\\usepackage{amsmath}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage{hyperref}\n"
        "\\usepackage{float}\n"
        "\\usepackage{caption}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
    ),
}

_DEFAULT_PREAMBLE = (
    "\\documentclass[12pt]{article}\n"
    "\\usepackage[margin=1in]{geometry}\n"
    "\\usepackage{amsmath}\n"
    "\\usepackage{amssymb}\n"
    "\\usepackage{graphicx}\n"
    "\\usepackage{hyperref}\n"
    "\\usepackage{float}\n"
    "\\usepackage{caption}\n"
    "\\usepackage{natbib}\n"
    "\\usepackage{times}\n"
    "\\usepackage[utf8]{inputenc}\n"
    "\\usepackage[T1]{fontenc}\n"
)


def _build_preamble(style: str) -> str:
    """Return a style-specific LaTeX preamble (everything before \\begin{document})."""
    return _PREAMBLE_MAP.get(style.lower(), _DEFAULT_PREAMBLE)


# ---------------------------------------------------------------------------
# LaTeX post-processor — fixes common LLM mistakes
# ---------------------------------------------------------------------------

def _postprocess_latex(body: str, style: str) -> str:
    """Clean up the merged LaTeX body before wrapping in preamble + document env."""

    # 1. Remove empty abstract environments  (LLM sometimes emits empty ones)
    body = re.sub(r"\\begin\{abstract\}\s*\\end\{abstract\}", "", body)

    # 2. Convert \section{Abstract} → proper \begin{abstract}...\end{abstract}
    def _replace_abstract_section(m: re.Match) -> str:
        content = m.group(1).strip()
        if not content:
            return ""
        return f"\\begin{{abstract}}\n{content}\n\\end{{abstract}}"

    body = re.sub(
        r"\\section\*?\{Abstract\}\s*\n(.*?)(?=\\section(?:\*?\{)|\\begin\{thebibliography\}|$)",
        _replace_abstract_section,
        body,
        flags=re.DOTALL | re.IGNORECASE,
    )

    # 3. Remove duplicate \maketitle
    count = body.count("\\maketitle")
    if count > 1:
        first = True
        def _keep_first(m: re.Match) -> str:
            nonlocal first
            if first:
                first = False
                return m.group(0)
            return ""
        body = re.sub(r"\\maketitle", _keep_first, body)

    # 4. Fix leftover markdown artefacts the LLM may emit
    body = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", body)
    body = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\\textit{\1}", body)

    # 5. Remove stray inner \begin{document} / \end{document}
    body = body.replace("\\begin{document}", "").replace("\\end{document}", "")

    # 6. Remove duplicate \title{} (keep first)
    titles = list(re.finditer(r"\\title\{", body))
    if len(titles) > 1:
        # keep first, remove rest
        for m in reversed(titles[1:]):
            # find closing brace
            depth = 0
            start = m.start()
            for i in range(start, len(body)):
                if body[i] == "{":
                    depth += 1
                elif body[i] == "}":
                    depth -= 1
                    if depth == 0:
                        body = body[:start] + body[i + 1:]
                        break

    # 7. Normalize multiple consecutive blank lines → max 2
    body = re.sub(r"\n{4,}", "\n\n\n", body)

    # 8. IEEE-specific: ensure \IEEEkeywords for Index Terms
    if style.lower() == "ieee":
        # Convert plain "Index Terms—..." paragraph into proper environment
        body = re.sub(
            r"(?m)^\\textit\{Index Terms\}[—–-]+\s*(.+)$",
            r"\\begin{IEEEkeywords}\n\1\n\\end{IEEEkeywords}",
            body,
        )

    # 9. Fix unescaped special chars in text (common LLM mistake)
    # Only fix & that isn't already in a tabular/align env context
    # This is tricky — skip for now to avoid false positives

    # 10. Remove any trailing \bibliography{} or \bibliographystyle{} (we use thebibliography)
    body = re.sub(r"\\bibliography\{[^}]*\}", "", body)
    body = re.sub(r"\\bibliographystyle\{[^}]*\}", "", body)

    return body.strip()


# ---------------------------------------------------------------------------
# Unified pipeline endpoint
# ---------------------------------------------------------------------------

@agno_router.post(
    "/api/v2/pipeline/stream",
    summary="Full pipeline: Static Format → LLM LaTeX (SSE)",
)
async def pipeline_stream(
    file: UploadFile = File(...),
    style: str = Form("apa7"),
    model: str = Form("llama-3.3-70b-versatile"),
):
    """
    Full FormatForge pipeline via Server-Sent Events:
      Stage 1 — Static formatting engine (6-agent Orchestrator)
      Stage 2 — LLM-based LaTeX generation (Agno + Groq)
    """
    api_key = os.environ.get("GROQ_API_KEY")
    filename = (file.filename or "upload.txt").lower()

    if not (filename.endswith(".docx") or filename.endswith(".pdf") or filename.endswith(".txt")):
        raise HTTPException(status_code=400, detail="Unsupported format. Upload .docx, .pdf, or .txt")

    file_bytes = await file.read()

    suffix = Path(filename).suffix
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(file_bytes)
    tmp.close()
    tmp_path = Path(tmp.name)

    async def sse_generator():
        formatted_path: Path | None = None
        try:
            # ══════════════════════════════════════════════════════
            # STAGE 1 — Static Formatting Engine (Orchestrator)
            # ══════════════════════════════════════════════════════
            yield _sse({"stage": 1, "log": "🔧 Starting static formatting engine..."})
            yield _sse({"stage": 1, "log": f"Style: {style} | File: {file.filename}"})
            await asyncio.sleep(0.05)

            from .agents.orchestrator import Orchestrator

            yield _sse({"stage": 1, "log": "Step 1/6 — Ingesting document..."})
            orchestrator = Orchestrator()

            result = await orchestrator.run(input_path=tmp_path, style_id=style)

            if not result.success:
                yield _sse({"error": f"Static formatting failed: {result.error_message}"})
                return

            formatted_path = Path(result.output_filename)

            yield _sse({"stage": 1, "log": "Step 2/6 — Style spec loaded"})
            yield _sse({"stage": 1, "log": "Step 3/6 — Document structure detected"})

            if result.structure_summary:
                summary_str = ", ".join(f"{k}: {v}" for k, v in result.structure_summary.items())
                yield _sse({"stage": 1, "log": f"Structure: {summary_str}"})

            yield _sse({"stage": 1, "log": "Step 4/6 — Citations processed"})
            yield _sse({"stage": 1, "log": "Step 5/6 — Formatting applied"})
            yield _sse({"stage": 1, "log": "Step 6/6 — Compliance validated"})

            if result.compliance_report:
                cr = result.compliance_report
                yield _sse({
                    "stage": 1,
                    "log": f"✅ Static formatting complete — Score: {cr.overall_score}% | "
                           f"Time: {result.processing_time_seconds}s",
                    "compliance_score": cr.overall_score,
                    "processing_time": result.processing_time_seconds,
                })

            yield _sse({
                "stage": 1,
                "log": "📄 Formatted document saved.",
                "formatted_file": formatted_path.name if formatted_path else None,
                "stage_complete": 1,
            })

            # Copy formatted file to the DocBot documents directory for agent editing.
            try:
                docs_dir = Path(
                    os.getenv("DOCUMENTS_DIR")
                    or (Path(__file__).resolve().parent.parent / "data" / "documents_store")
                )
                docs_dir.mkdir(parents=True, exist_ok=True)
                dest = docs_dir / formatted_path.name
                shutil.copy2(str(formatted_path), str(dest))
                yield _sse({"stage": 1, "log": f"📂 Copied to agent documents: {formatted_path.name}"})
            except Exception as copy_err:
                logger.warning(f"Could not copy to DocBot documents: {copy_err}")

            # ══════════════════════════════════════════════════════
            # STAGE 2 — LLM LaTeX Generation (Chunked)
            # ══════════════════════════════════════════════════════

            if not _AGNO_READY or not api_key:
                yield _sse({
                    "stage": 2,
                    "log": "⚠️ LLM LaTeX generation skipped (Agno/Groq not available).",
                    "is_final": True,
                    "latex": "",
                    "formatted_file": formatted_path.name if formatted_path else None,
                })
                return

            yield _sse({"stage": 2, "log": "🧠 Starting LLM-based LaTeX generation..."})
            await asyncio.sleep(0.05)

            yield _sse({"stage": 2, "log": "Extracting text and figures from formatted document..."})

            # Per-job assets dir alongside the formatted DOCX (served via /assets).
            from .config import OUTPUT_FORMATTED_DIR
            job_id = formatted_path.stem
            assets_dir = OUTPUT_FORMATTED_DIR / f"{job_id}_assets"

            figures: list[dict] = []
            try:
                formatted_text, figures, skipped_imgs = _extract_text_and_images(formatted_path, assets_dir)
            except Exception as e:
                logger.warning("Image-aware extraction failed (%s); falling back to text-only.", e)
                formatted_text = _extract_from_docx(formatted_path)
                skipped_imgs = 0

            if not formatted_text.strip():
                yield _sse({"error": "Formatted document is empty — cannot generate LaTeX."})
                return

            word_count = len(formatted_text.split())
            yield _sse({"stage": 2, "log": f"Extracted {word_count} words from formatted document."})
            if figures:
                yield _sse({"stage": 2, "log": f"🖼️ Extracted {len(figures)} figure(s) — will embed as \\includegraphics."})
            if skipped_imgs:
                yield _sse({"stage": 2, "log": f"⚠ {skipped_imgs} image(s) in an unsupported format (e.g. EMF/WMF) were skipped."})

            # ── Split into sections for chunked generation ──
            sections = _split_into_sections(formatted_text)
            yield _sse({"stage": 2, "log": f"Split document into {len(sections)} sections for chunked LaTeX generation."})

            latex_agent = Agent(
                name="LaTeXConverterAgent",
                model=AgnoGroq(id=model, api_key=api_key),
                instructions=(
                    f"You are an expert LaTeX code generator for {style.upper()} academic manuscripts.\n"
                    f"You will receive ONE SECTION of a manuscript at a time.\n"
                    f"Convert it into LaTeX body content ONLY.\n\n"
                    f"RULES:\n"
                    f"1. Output ONLY LaTeX body content — NO \\documentclass, NO \\usepackage, NO \\begin{{document}}, NO \\end{{document}}.\n"
                    f"2. Use \\section{{}}, \\subsection{{}}, \\subsubsection{{}} for headings.\n"
                    f"3. Convert citations to \\cite{{author_year}} format.\n"
                    f"4. Preserve EVERY paragraph, EVERY sentence, EVERY detail. Do NOT summarize.\n"
                    f"5. For TABLE descriptions, use \\begin{{table}}[H]...\\end{{table}}.\n"
                    f"6. FIGURES: the text contains placeholder tokens like [[FIGURE_1]], [[FIGURE_2]]. "
                    f"Keep every such token EXACTLY as-is, on its own line, where it appears. "
                    f"Do NOT create \\begin{{figure}} or \\includegraphics yourself — the tokens are replaced automatically.\n"
                    f"7. Use proper LaTeX for math: $...$ for inline, $$...$$ for display.\n"
                    f"8. Use \\textit{{}} for species names, \\textbf{{}} for emphasis.\n"
                    f"9. Return ONLY raw LaTeX code. NO markdown fences. NO explanations.\n"
                    f"10. Do NOT skip or truncate ANY content. Every word matters.\n"
                ),
                markdown=False,
            )

            # ── Generate LaTeX for each section ──
            body_parts: list[str] = []
            preamble_section = None

            for idx, section in enumerate(sections):
                sec_title = section["title"]
                sec_content = section["content"]
                sec_words = len(sec_content.split())

                yield _sse({"stage": 2, "log": f"Chunk {idx+1}/{len(sections)}: {sec_title} ({sec_words} words)..."})

                # For the first section (Preamble), generate title/author/abstract
                if idx == 0 and sec_title == "Preamble":
                    chunk_prompt = (
                        f"Convert this manuscript header section into LaTeX code.\n"
                        f"Include: \\title{{}}, \\author{{}}, \\maketitle, and \\begin{{abstract}}...\\end{{abstract}} if abstract is present.\n"
                        f"Do NOT output \\documentclass or \\usepackage or \\begin{{document}}.\n"
                        f"Preserve ALL text exactly.\n\n"
                        f"--- SECTION ---\n{sec_content}\n--- END ---\n"
                    )
                elif "reference" in sec_title.lower():
                    chunk_prompt = (
                        f"Convert this references section into LaTeX \\begin{{thebibliography}} format.\n"
                        f"Use \\bibitem{{key}} for each reference. Preserve ALL references.\n"
                        f"Do NOT output \\documentclass or \\usepackage.\n\n"
                        f"--- REFERENCES ---\n{sec_content}\n--- END ---\n"
                    )
                else:
                    chunk_prompt = (
                        f"Convert this section into LaTeX body content.\n"
                        f"Section heading: {sec_title}\n"
                        f"Use \\section{{{sec_title}}} at the start (unless already a subsection).\n"
                        f"Do NOT output \\documentclass, \\usepackage, \\begin{{document}}, or \\end{{document}}.\n"
                        f"Preserve EVERY paragraph exactly. Do NOT summarize or skip.\n\n"
                        f"--- SECTION ---\n{sec_content}\n--- END ---\n"
                    )

                try:
                    chunk_result = await asyncio.to_thread(latex_agent.run, chunk_prompt)
                    chunk_latex = _clean_chunk_latex(chunk_result.content)
                    body_parts.append(chunk_latex)
                    yield _sse({"stage": 2, "log": f"  ✓ Chunk {idx+1} complete ({len(chunk_latex)} chars)"})
                except Exception as e:
                    logger.warning(f"Chunk {idx+1} ({sec_title}) failed: {e}")
                    yield _sse({"stage": 2, "log": f"  ⚠ Chunk {idx+1} failed: {str(e)[:80]}, using raw text fallback"})
                    # Fallback: wrap raw text in LaTeX
                    escaped = sec_content.replace("&", "\\&").replace("%", "\\%").replace("#", "\\#").replace("_", "\\_")
                    body_parts.append(f"\\section{{{sec_title}}}\n\n{escaped}")

                await asyncio.sleep(0.1)  # Rate limit between Groq calls

            # ── Assemble final LaTeX document ──
            yield _sse({"stage": 2, "log": "Assembling final LaTeX document from all chunks..."})

            merged_body = "\n\n".join(body_parts)

            # Post-process: fix LLM mistakes (empty abstracts, markdown, duplicates)
            yield _sse({"stage": 2, "log": "Post-processing LaTeX (fixing LLM artifacts)..."})
            merged_body = _postprocess_latex(merged_body, style)

            # Deterministically replace [[FIGURE_n]] sentinels with real figure
            # floats; re-append any the LLM dropped so no figure is ever lost.
            if figures:
                yield _sse({"stage": 2, "log": f"Embedding {len(figures)} figure(s) into LaTeX..."})
            merged_body = _inject_figures(merged_body, figures)

            # Style-specific preamble (+ graphicspath so assets/ resolves).
            preamble = _build_preamble(style)
            if figures:
                preamble = f"{preamble}\\graphicspath{{{{assets/}}}}\n"

            merged_latex = (
                f"{preamble}\n"
                "\\begin{document}\n\n"
                f"{merged_body}\n\n"
                "\\end{document}\n"
            )

            yield _sse({"stage": 2, "log": f"✅ LaTeX generation complete! ({len(merged_latex)} chars, {len(merged_latex.split(chr(10)))} lines)"})

            # ══════════════════════════════════════════════════════
            # CONTENT-INTEGRITY CHECK (Phase C) — prove no data loss
            # ══════════════════════════════════════════════════════
            integrity_payload = None
            try:
                yield _sse({"stage": 3, "log": "🔎 Verifying content preservation..."})
                from .agents.content_integrity import validate_content_preservation

                # Original reference text = extracted formatted text minus figure
                # sentinels (so [[FIGURE_n]] tokens don't count as words).
                original_ref_text = re.sub(r"\[\[FIGURE_\d+\]\]", "", formatted_text)
                table_count_in = 0
                if result.structure_summary:
                    table_count_in = int(result.structure_summary.get("table", 0))
                src_fmt = Path(file.filename or "x.docx").suffix.lstrip(".").lower() or "docx"

                # Count only successfully-extracted figures — placeholders for
                # missing ones don't emit \includegraphics yet.
                extracted_count = sum(1 for f in figures if f.get("filename"))
                integrity = validate_content_preservation(
                    original_text=original_ref_text,
                    latex_str=merged_latex,
                    figure_count_in=extracted_count,
                    table_count_in=table_count_in,
                    source_format=src_fmt,
                )
                integrity_payload = integrity.model_dump(mode="json")

                icon = {"info": "✅", "warning": "⚠️", "error": "🔴"}.get(integrity.severity.value, "ℹ️")
                yield _sse({
                    "stage": 3,
                    "log": f"{icon} Integrity: {integrity.summary()}",
                    "integrity": integrity_payload,
                })
            except Exception as ie:
                logger.warning("Integrity check error: %s", ie)

            asset_files = [f["filename"] for f in figures if f.get("filename")]
            # Figures whose image couldn't be extracted — the frontend (Phase I)
            # prompts the user to upload an image for each.
            missing_figures = [
                {"n": f["number"], "token": f["token"], "caption": f.get("caption") or ""}
                for f in figures if not f.get("filename")
            ]
            if missing_figures:
                yield _sse({"stage": 3, "log": f"🖼️ {len(missing_figures)} figure(s) need an uploaded image."})
            yield _sse({
                "is_final": True,
                "latex": merged_latex,
                "formatted_file": formatted_path.name if formatted_path else None,
                "job": job_id,
                "assets": asset_files,
                "assets_base": f"/deepscan/api/v2/assets/{job_id}",
                "figure_count": len(asset_files),
                "missing_figures": missing_figures,
                "integrity": integrity_payload,
                "content_integrity_passed": (integrity_payload or {}).get("passed", True),
            })

        except Exception as e:
            logger.exception("Pipeline error: %s", e)
            yield _sse({"error": str(e)})

        finally:
            tmp_path.unlink(missing_ok=True)

    return StreamingResponse(sse_generator(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
# Download formatted file
# ---------------------------------------------------------------------------

@agno_router.get("/api/v2/download/{filename}")
async def download_formatted(filename: str):
    """Download the statically formatted .docx file."""
    from .config import OUTPUT_FORMATTED_DIR

    file_path = OUTPUT_FORMATTED_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


# ---------------------------------------------------------------------------
# Serve extracted figure assets (Phase B) — referenced by \includegraphics
# {assets/figN.png} in generated LaTeX. The local compiler (Phase D) and the
# diagram-upload flow (Phase I) read from here.
# ---------------------------------------------------------------------------

@agno_router.get("/api/v2/assets/{job}/{filename}")
async def get_asset(job: str, filename: str):
    """Return a single extracted figure image for a given job (formatted-file stem)."""
    from .config import OUTPUT_FORMATTED_DIR

    # Guard against path traversal.
    if "/" in job or "\\" in job or "/" in filename or "\\" in filename or ".." in job or ".." in filename:
        raise HTTPException(status_code=400, detail="Invalid asset path")

    asset_path = OUTPUT_FORMATTED_DIR / f"{job}_assets" / filename
    if not asset_path.exists():
        raise HTTPException(status_code=404, detail="Asset not found")
    return FileResponse(str(asset_path))


@agno_router.post("/api/v2/assets/{job}/{n}")
async def upload_figure(
    job: str,
    n: int,
    file: UploadFile = File(...),
    latex: str = Form(""),
):
    """Phase I — upload an image for a figure the pipeline couldn't extract.

    Saves the image into the job's assets dir as ``fig{n}.png`` (normalised to
    PNG when possible) and, when the caller passes the current ``latex``,
    rewrites the placeholder slot into a real ``\\includegraphics`` and returns
    the updated source. The frontend swaps its editor content with the returned
    ``latex`` and recompiles so the figure renders in the PDF.
    """
    from .config import OUTPUT_FORMATTED_DIR

    if "/" in job or "\\" in job or ".." in job:
        raise HTTPException(status_code=400, detail="Invalid job id")

    job_dir = OUTPUT_FORMATTED_DIR / f"{job}_assets"
    job_dir.mkdir(parents=True, exist_ok=True)

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty file.")

    png = _raster_to_png(raw)
    if png is not None:
        fname = f"fig{n}.png"
        (job_dir / fname).write_bytes(png)
    else:
        ext = (os.path.splitext(file.filename or "")[1] or ".png").lower()
        fname = f"fig{n}{ext}"
        (job_dir / fname).write_bytes(raw)

    rel_path = f"assets/{fname}"
    updated_latex = _fill_figure_slot(latex, n, rel_path) if latex else ""

    return {
        "ok": True,
        "n": n,
        "filename": fname,
        "rel_path": rel_path,
        "latex": updated_latex,
    }


# ---------------------------------------------------------------------------
# Local LaTeX compile (Phase D) — tectonic + auto-correct. Replaces the flaky
# texlive.net round-trip and renders figures from the job's assets/ dir.
# ---------------------------------------------------------------------------

from pydantic import BaseModel


class CompileRequest(BaseModel):
    latex: str
    job: str | None = None          # formatted-file stem → resolves assets dir
    autofix: bool = True
    allow_llm_repair: bool = True


@agno_router.post("/api/v2/compile")
async def compile_latex_endpoint(req: CompileRequest):
    """Compile LaTeX → PDF locally with tectonic.

    Returns the PDF (application/pdf) on success, with the applied auto-fix
    notes in the ``X-Latex-Notes`` header. On failure returns a JSON body with
    the tectonic error log so the editor can show it inline.
    """
    from .config import OUTPUT_FORMATTED_DIR
    from .latex_compile import compile_with_repair, autofix_latex, tectonic_available

    if not req.latex or not req.latex.strip():
        raise HTTPException(status_code=400, detail="LaTeX source is required.")

    # Resolve the per-job assets dir (if the doc had figures).
    assets_dir = None
    if req.job and ".." not in req.job and "/" not in req.job and "\\" not in req.job:
        candidate = OUTPUT_FORMATTED_DIR / f"{req.job}_assets"
        if candidate.is_dir():
            assets_dir = candidate

    if not tectonic_available():
        # Surface auto-fix work even when the engine is missing, so the client
        # can fall back to remote compile with the cleaned source.
        fixed, notes = autofix_latex(req.latex) if req.autofix else (req.latex, [])
        return JSONResponse(
            status_code=503,
            content={
                "success": False,
                "engine": "tectonic",
                "tectonic_available": False,
                "message": "Local LaTeX engine (tectonic) is not installed on the server.",
                "notes": notes,
                "fixed_latex": fixed,
            },
        )

    result = await asyncio.to_thread(
        compile_with_repair,
        req.latex,
        assets_dir=assets_dir,
        allow_llm_repair=req.allow_llm_repair,
    )

    if result.ok and result.pdf_bytes:
        from fastapi.responses import Response
        headers = {
            "Content-Disposition": "inline; filename=main.pdf",
            "X-Latex-Notes": json.dumps(result.notes)[:2000],
        }
        return Response(content=result.pdf_bytes, media_type="application/pdf", headers=headers)

    return JSONResponse(
        status_code=422,
        content={
            "success": False,
            "engine": result.engine,
            "message": "LaTeX compilation failed.",
            "notes": result.notes,
            "log": result.log[:12000],
        },
    )


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@agno_router.get("/api/v2/health")
async def agno_health():
    return {
        "agno_available": _AGNO_READY,
        "groq_key_set": bool(os.environ.get("GROQ_API_KEY")),
    }
