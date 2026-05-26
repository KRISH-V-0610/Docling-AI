"""
FormatForge AI — Agent 1: Ingest & Parse
Converts DOCX / PDF / TXT → DocIR JSON.

Phase 1 — Full DOCX parser with:
 • Interleaved paragraph / table ordering (XML body walk)
 • Inherited-font resolution via style chain
 • Inline-image / figure detection
 • Rich table-cell formatting extraction
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from ..schemas.docir import (
    DocElement,
    DocIR,
    DocMetadata,
    ElementRole,
    ElementType,
    ParagraphFormatting,
    RunFormatting,
    TableCell,
    TableData,
)

logger = logging.getLogger(__name__)


class IngestAgent:
    """Agent 1 — Parse manuscript files into DocIR."""

    def parse(self, file_path: Path) -> DocIR:
        """
        Parse a manuscript file into DocIR.

        Args:
            file_path: Path to the manuscript file.

        Returns:
            DocIR with raw elements (roles NOT yet assigned — that's Agent 3).
        """
        suffix = file_path.suffix.lower()
        if suffix == ".docx":
            return self._parse_docx(file_path)
        elif suffix == ".pdf":
            return self._parse_pdf(file_path)
        elif suffix == ".txt":
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    # ── DOCX Parser ──────────────────────────────────────────

    def _parse_docx(self, file_path: Path) -> DocIR:
        """Parse a DOCX file into DocIR using python-docx with interleaved ordering."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document(str(file_path))
        elements: list[DocElement] = []
        table_count = 0
        figure_count = 0
        elem_idx = 0

        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }

        # Resolve the document's default font for fallback
        default_font_name: Optional[str] = None
        default_font_size_pt: Optional[float] = None
        try:
            normal_style = doc.styles["Normal"]
            if normal_style.font.name:
                default_font_name = normal_style.font.name
            if normal_style.font.size:
                default_font_size_pt = round(normal_style.font.size / 12700, 1)
        except Exception:
            pass

        # ── Build an ordering map from XML body ──────────────
        # This ensures paragraphs & tables appear in true document order.
        body_xml = doc.element.body
        paragraphs = doc.paragraphs
        tables = doc.tables
        p_idx = 0
        t_idx = 0
        order: list[tuple[str, int]] = []  # ('p', idx) | ('t', idx)

        for child in body_xml:
            if child.tag == qn("w:p"):
                order.append(("p", p_idx))
                p_idx += 1
            elif child.tag == qn("w:tbl"):
                order.append(("t", t_idx))
                t_idx += 1

        # ── Iterate in document order ────────────────────────
        for item_type, item_idx in order:
            elem_idx += 1
            elem_id = f"elem_{elem_idx:04d}"

            if item_type == "p" and item_idx < len(paragraphs):
                para = paragraphs[item_idx]

                # Check for inline images / drawings
                para_xml = para._element
                drawings = para_xml.findall(f".//{qn('w:drawing')}")
                has_image = len(drawings) > 0
                if has_image:
                    figure_count += len(drawings)

                element = self._parse_paragraph(
                    para, elem_id, alignment_map,
                    default_font_name, default_font_size_pt,
                )
                if has_image:
                    element.figure_number = figure_count
                elements.append(element)

            elif item_type == "t" and item_idx < len(tables):
                table_count += 1
                element = self._parse_table(tables[item_idx], elem_id, table_count)
                elements.append(element)

        # ── Build metadata ───────────────────────────────────
        metadata = DocMetadata(
            source_filename=file_path.name,
            source_format="docx",
            total_paragraphs=sum(1 for e in elements if e.type == ElementType.PARAGRAPH),
            total_tables=table_count,
            total_figures=figure_count,
        )

        docir = DocIR(metadata=metadata, elements=elements)
        logger.info(
            "Ingested %s — %d elements (%d paragraphs, %d tables, %d figures)",
            file_path.name,
            len(elements),
            metadata.total_paragraphs,
            metadata.total_tables,
            metadata.total_figures,
        )
        return docir

    # ── Paragraph extractor ──────────────────────────────────

    def _parse_paragraph(
        self,
        para,
        elem_id: str,
        alignment_map: dict,
        default_font_name: Optional[str],
        default_font_size_pt: Optional[float],
    ) -> DocElement:
        """Extract a single paragraph into a DocElement with inherited-font resolution."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F811

        pf = para.paragraph_format
        alignment_val = alignment_map.get(para.alignment, None) if para.alignment else None

        # Line spacing
        line_spacing = None
        line_spacing_rule = None
        if pf.line_spacing is not None:
            line_spacing = float(pf.line_spacing)
        if pf.line_spacing_rule is not None:
            line_spacing_rule = str(pf.line_spacing_rule)

        # Indentation  (EMU → inches)
        first_line_indent = None
        if pf.first_line_indent is not None:
            first_line_indent = round(pf.first_line_indent / 914400, 3)

        left_indent = None
        if pf.left_indent is not None:
            left_indent = round(pf.left_indent / 914400, 3)

        # Space before / after (EMU → pt)
        space_before = None
        if pf.space_before is not None:
            space_before = round(pf.space_before / 12700, 1)

        space_after = None
        if pf.space_after is not None:
            space_after = round(pf.space_after / 12700, 1)

        # ── Parse runs with inherited-font resolution ────────
        runs_data: list[RunFormatting] = []
        dominant_font = None
        dominant_size = None
        is_bold = None
        is_italic = None

        # Style-level font fallback
        style_font_name: Optional[str] = None
        style_font_size_pt: Optional[float] = None
        try:
            if para.style and para.style.font:
                style_font_name = para.style.font.name
                if para.style.font.size:
                    style_font_size_pt = round(para.style.font.size / 12700, 1)
        except Exception:
            pass

        for run in para.runs:
            font = run.font

            # Resolve font name: run → style → doc default
            resolved_font = font.name or style_font_name or default_font_name
            # Resolve font size: run → style → doc default
            resolved_size: Optional[float] = None
            if font.size:
                resolved_size = round(font.size / 12700, 1)
            elif style_font_size_pt:
                resolved_size = style_font_size_pt
            elif default_font_size_pt:
                resolved_size = default_font_size_pt

            rf = RunFormatting(
                text=run.text,
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
                font_name=resolved_font,
                font_size_pt=resolved_size,
                superscript=font.superscript,
                subscript=font.subscript,
            )
            runs_data.append(rf)

            # Track dominant formatting (first text-bearing run wins)
            if run.text.strip():
                if dominant_font is None:
                    dominant_font = resolved_font
                if dominant_size is None:
                    dominant_size = resolved_size
                if is_bold is None:
                    is_bold = run.bold
                if is_italic is None:
                    is_italic = run.italic

        formatting = ParagraphFormatting(
            font_name=dominant_font,
            font_size_pt=dominant_size,
            bold=is_bold,
            italic=is_italic,
            alignment=alignment_val,
            line_spacing=line_spacing,
            line_spacing_rule=line_spacing_rule,
            space_before_pt=space_before,
            space_after_pt=space_after,
            first_line_indent_inches=first_line_indent,
            left_indent_inches=left_indent,
        )

        return DocElement(
            id=elem_id,
            type=ElementType.PARAGRAPH,
            role=ElementRole.UNKNOWN,  # Agent 3 will assign roles
            content=para.text,
            original_style_name=para.style.name if para.style else None,
            formatting=formatting,
            runs=runs_data,
        )

    # ── Table extractor ──────────────────────────────────────

    def _parse_table(self, table, elem_id: str, table_count: int) -> DocElement:
        """Extract a table into a DocElement with cell-level formatting."""
        rows_data: list[list[TableCell]] = []
        for r_idx, row in enumerate(table.rows):
            row_cells: list[TableCell] = []
            for c_idx, cell in enumerate(row.cells):
                cell_bold = None
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell_bold = cell.paragraphs[0].runs[0].bold
                row_cells.append(TableCell(
                    text=cell.text.strip(),
                    row=r_idx,
                    col=c_idx,
                    bold=cell_bold,
                ))
            rows_data.append(row_cells)

        num_cols = 0
        try:
            num_cols = len(table.columns)
        except Exception:
            if rows_data:
                num_cols = len(rows_data[0])

        td = TableData(
            rows=rows_data,
            num_rows=len(table.rows),
            num_cols=num_cols,
        )

        return DocElement(
            id=elem_id,
            type=ElementType.TABLE,
            role=ElementRole.TABLE,
            content="",
            table_data=td,
            table_number=table_count,
        )

    # ── PDF Parser (stub) ────────────────────────────────────

    def _parse_pdf(self, file_path: Path) -> DocIR:
        """Parse a PDF file into DocIR (basic extraction with PyMuPDF)."""
        logger.warning("PDF parsing is a basic stub — use DOCX for best results.")
        raise NotImplementedError("PDF parsing not yet implemented. Please upload a .docx file.")

    # ── Plain text Parser (stub) ─────────────────────────────

    def _parse_txt(self, file_path: Path) -> DocIR:
        """Parse a plain text file into DocIR."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            elements: list[DocElement] = []
            
            # Split by double newline to approximate paragraphs
            paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
            
            for idx, text in enumerate(paragraphs):
                elem_id = f"txt_elem_{idx:04d}"
                # Create default formatting
                formatting = ParagraphFormatting(
                    font_name=None,
                    font_size_pt=None,
                    bold=False,
                    italic=False,
                    alignment="left"
                )
                
                element = DocElement(
                    id=elem_id,
                    type=ElementType.PARAGRAPH,
                    role=ElementRole.UNKNOWN,
                    content=text,
                    formatting=formatting,
                    runs=[RunFormatting(text=text, bold=False, italic=False)]
                )
                elements.append(element)
                
            metadata = DocMetadata(
                source_filename=file_path.name,
                source_format="txt",
                total_paragraphs=len(elements),
                total_tables=0,
                total_figures=0,
            )
            
            docir = DocIR(metadata=metadata, elements=elements)
            logger.info("Ingested %s — %d elements from TXT", file_path.name, len(elements))
            return docir
            
        except Exception as e:
            logger.error("Failed to parse TXT file: %s", str(e))
            raise ValueError(f"Failed to parse TXT: {str(e)}")
